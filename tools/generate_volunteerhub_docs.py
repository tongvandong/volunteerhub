from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "Tài liệu CNPM" / "generated"
DIAGRAM_DIR = OUT_DIR / "diagrams"


ACCENT = RGBColor(46, 116, 181)
ACCENT_DARK = RGBColor(31, 77, 120)
INK = RGBColor(31, 31, 31)
MUTED = RGBColor(95, 95, 95)
TABLE_HEADER_FILL = "F2F4F7"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    mar = tbl_pr.find(qn("w:tblCellMar"))
    if mar is None:
        mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360, indent_dxa: int = 120) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    set_cell_margins(table)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def apply_report_numbering(doc: Document) -> None:
    """Apply report/template-style numbering: I., 1.1, a)."""
    roman = [
        "I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X",
        "XI", "XII", "XIII", "XIV", "XV", "XVI", "XVII", "XVIII", "XIX", "XX",
    ]
    skip_h1 = {"NỘI DUNG SỬA ĐỔI", "QUÁ TRÌNH THAY ĐỔI", "TRANG KÝ", "MỤC LỤC"}
    start_numbering = False
    h1_count = 0
    h2_count = 0
    h3_count = 0

    for paragraph in doc.paragraphs:
        style = paragraph.style.name if paragraph.style is not None else ""
        text = paragraph.text.strip()
        if not text:
            continue

        if style == "Heading 1":
            if text in skip_h1 and not start_numbering:
                continue
            start_numbering = True
            h1_count += 1
            h2_count = 0
            h3_count = 0
            label = roman[h1_count - 1] if h1_count <= len(roman) else str(h1_count)
            paragraph.text = f"{label}. {text}"
        elif style == "Heading 2" and start_numbering:
            h2_count += 1
            h3_count = 0
            paragraph.text = f"{h1_count}.{h2_count} {text}"
        elif style == "Heading 3" and start_numbering:
            label = chr(ord("a") + (h3_count % 26))
            h3_count += 1
            paragraph.text = f"{label}) {text}"


def setup_document(title: str, subtitle: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, ACCENT, 16, 8),
        ("Heading 2", 13, ACCENT, 12, 6),
        ("Heading 3", 12, ACCENT_DARK, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ["List Bullet", "List Number"]:
        style = styles[list_style]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.167

    for sec in doc.sections:
        header = sec.header.paragraphs[0]
        header.text = title
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header.runs[0].font.size = Pt(9)
        header.runs[0].font.color.rgb = MUTED

        footer = sec.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer.add_run("VolunteerHub - CNPM")
        footer.runs[0].font.size = Pt(9)
        footer.runs[0].font.color.rgb = MUTED

    add_cover(doc, title, subtitle)
    return doc


def add_cover(doc: Document, title: str, subtitle: str) -> None:
    org_lines = [
        "Khoa Công nghệ thông tin - Học viện Kỹ thuật quân sự",
        "Trụ sở: 236 Hoàng Quốc Việt, Cầu Giấy, Hà Nội",
        "Email: fit@mta.edu.vn - Website: https://fit.mta.edu.vn",
    ]
    for line in org_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(line)
        r.font.size = Pt(10)
        r.font.color.rgb = MUTED
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("BÁO CÁO ĐỒ ÁN XÂY DỰNG PHẦN MỀM")
    r.bold = True
    r.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(34)
    r = p.add_run("HỆ THỐNG VOLUNTEERHUB")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = ACCENT_DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(36)
    r = p.add_run(subtitle)
    r.italic = True
    r.font.size = Pt(12)
    r.font.color.rgb = MUTED

    metadata = [
        ("Mã dự án", "VolunteerHub"),
        ("Mã tài liệu", "VH-SRS-1.0" if "YÊU CẦU" in title else "VH-DD-1.0"),
        ("Tên hệ thống", "VolunteerHub - Cổng sự kiện tình nguyện"),
        ("Loại tài liệu", title),
        ("Phiên bản", "1.0"),
        ("Ngày lập", "11/06/2026"),
    ]
    add_table(doc, ["Thuộc tính", "Giá trị"], metadata, [2200, 7160])

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Hà Nội, 06/2026")
    r.font.size = Pt(11)
    r.font.color.rgb = MUTED
    doc.add_page_break()


def add_history(doc: Document) -> None:
    doc.add_heading("QUÁ TRÌNH THAY ĐỔI", level=1)
    rows = [
        ("11/06/2026", "Tạo tài liệu", "M", "Tạo mới tài liệu theo template CNPM.", "Nhóm thực hiện", "1"),
    ]
    add_table(
        doc,
        ["Ngày", "Mục sửa đổi", "M* / S, X", "Nội dung sửa đổi", "Người sửa đổi", "Lần sửa đổi"],
        rows,
        [1300, 1900, 1000, 3000, 1500, 660],
    )
    p = doc.add_paragraph("*M: Mới, S: Sửa, X: Xóa")
    p.runs[0].italic = True


def add_signature_page(doc: Document) -> None:
    doc.add_heading("TRANG KÝ", level=1)
    for role in ["NGƯỜI LẬP", "NGƯỜI KIỂM TRA", "NGƯỜI PHÊ DUYỆT"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{role}:")
        r.bold = True
        p.add_run("\t\t........................................\tNgày: ....../....../......")
        p2 = doc.add_paragraph("<Vị trí>")
        p2.paragraph_format.left_indent = Inches(1.25)
        p2.runs[0].italic = True
        p2.runs[0].font.color.rgb = MUTED


def add_toc_note(doc: Document) -> None:
    doc.add_heading("MỤC LỤC", level=1)
    p = doc.add_paragraph(
        "Mục lục được cập nhật tự động trong Microsoft Word bằng thao tác References > Update Table sau khi mở tài liệu."
    )
    p.runs[0].italic = True
    for item in ["I. GIỚI THIỆU", "II. BỐI CẢNH VÀ ĐỊNH HƯỚNG XÂY DỰNG", "III. TỔNG QUAN HỆ THỐNG"]:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.25)
    doc.add_page_break()


def add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Iterable[Sequence[str]],
    widths: Sequence[int] | None = None,
) -> None:
    rows = list(rows)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = str(text)
        set_cell_shading(cell, TABLE_HEADER_FILL)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            set_cell_width(cell, widths[i])
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = str(text)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                set_cell_width(cells[i], widths[i])
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_pseudocode(doc: Document, code: str) -> None:
    for line in code.strip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r.font.size = Pt(9)
    doc.add_paragraph()


def _font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
        Path("C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def _wrap(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        trial = f"{current} {word}".strip()
        if draw.textbbox((0, 0), trial, font=font)[2] <= max_width or not current:
            current = trial
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def _center_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font, fill=(28, 39, 51)) -> None:
    x1, y1, x2, y2 = box
    lines = _wrap(draw, text, font, x2 - x1 - 18)
    line_h = font.size + 4
    y = y1 + max(4, ((y2 - y1) - line_h * len(lines)) // 2)
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        draw.text((x1 + ((x2 - x1) - (bbox[2] - bbox[0])) // 2, y), line, font=font, fill=fill)
        y += line_h


def _box(draw, xy, text, fill="#F7FAFC", outline="#2E74B5", radius=10, font=None, width=2) -> None:
    font = font or _font(18)
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)
    _center_text(draw, xy, text, font)


def _oval(draw, xy, text, fill="#FFFFFF", outline="#2E74B5", font=None, width=2) -> None:
    font = font or _font(17)
    draw.ellipse(xy, fill=fill, outline=outline, width=width)
    _center_text(draw, xy, text, font)


def _arrow(draw, start, end, fill="#34495E", width=3) -> None:
    draw.line([start, end], fill=fill, width=width)
    import math
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 12
    pts = [
        end,
        (end[0] - size * math.cos(angle - math.pi / 6), end[1] - size * math.sin(angle - math.pi / 6)),
        (end[0] - size * math.cos(angle + math.pi / 6), end[1] - size * math.sin(angle + math.pi / 6)),
    ]
    draw.polygon(pts, fill=fill)


def _actor(draw, x: int, y: int, label: str) -> None:
    pen = "#2C3E50"
    draw.ellipse((x - 14, y, x + 14, y + 28), outline=pen, width=3)
    draw.line((x, y + 28, x, y + 82), fill=pen, width=3)
    draw.line((x - 34, y + 48, x + 34, y + 48), fill=pen, width=3)
    draw.line((x, y + 82, x - 28, y + 122), fill=pen, width=3)
    draw.line((x, y + 82, x + 28, y + 122), fill=pen, width=3)
    _center_text(draw, (x - 70, y + 126, x + 70, y + 172), label, _font(16, True))


def _new_diagram(name: str, size=(1600, 1000)) -> tuple[Path, Image.Image, ImageDraw.ImageDraw]:
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    path = DIAGRAM_DIR / name
    img = Image.new("RGB", size, "white")
    draw = ImageDraw.Draw(img)
    return path, img, draw


def _title(draw, text: str, width: int) -> None:
    draw.text((40, 24), text, font=_font(32, True), fill="#1F4E79")
    draw.line((40, 70, width - 40, 70), fill="#D9E2F3", width=3)


def _dashed_line(draw, start, end, fill="#34495E", width=2, dash=12) -> None:
    import math
    x1, y1 = start
    x2, y2 = end
    length = math.hypot(x2 - x1, y2 - y1)
    if length == 0:
        return
    dx = (x2 - x1) / length
    dy = (y2 - y1) / length
    pos = 0
    while pos < length:
        seg_end = min(pos + dash, length)
        draw.line(
            [(x1 + dx * pos, y1 + dy * pos), (x1 + dx * seg_end, y1 + dy * seg_end)],
            fill=fill,
            width=width,
        )
        pos += dash * 2


def _dashed_dependency(draw, start, end, label: str = "<<include>>", fill="#1F4E79") -> None:
    import math
    _dashed_line(draw, start, end, fill=fill, width=3, dash=12)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 14
    p1 = (end[0] - size * math.cos(angle - math.pi / 6), end[1] - size * math.sin(angle - math.pi / 6))
    p2 = (end[0] - size * math.cos(angle + math.pi / 6), end[1] - size * math.sin(angle + math.pi / 6))
    draw.line([p1, end, p2], fill=fill, width=3)
    mx, my = (start[0] + end[0]) // 2, (start[1] + end[1]) // 2
    bbox = draw.textbbox((0, 0), label, font=_font(16, True))
    draw.rectangle((mx - 48, my - 18, mx + (bbox[2] - bbox[0]) - 42, my + 6), fill="white")
    draw.text((mx - 46, my - 18), label, font=_font(16, True), fill=fill)


def _return_arrow(draw, start, end, label: str) -> None:
    _dashed_line(draw, start, end)
    _arrow(draw, (end[0] + 1, end[1]), end, width=2)
    draw.text((min(start[0], end[0]) + 12, start[1] - 24), label, font=_font(15), fill="#2C3E50")


def _activation(draw, x: int, y1: int, y2: int) -> None:
    draw.rectangle((x - 8, y1, x + 8, y2), fill="#FFFFFF", outline="#34495E", width=2)


def _initial_node(draw, cx: int, cy: int) -> None:
    draw.ellipse((cx - 13, cy - 13, cx + 13, cy + 13), fill="#1F2933", outline="#1F2933")


def _final_node(draw, cx: int, cy: int) -> None:
    draw.ellipse((cx - 16, cy - 16, cx + 16, cy + 16), outline="#1F2933", width=3)
    draw.ellipse((cx - 9, cy - 9, cx + 9, cy + 9), fill="#1F2933", outline="#1F2933")


def _decision(draw, box, text: str, fill="#FFF7E6") -> None:
    x1, y1, x2, y2 = box
    draw.polygon([(x1 + x2) // 2, y1, x2, (y1 + y2) // 2, (x1 + x2) // 2, y2, x1, (y1 + y2) // 2], fill=fill, outline="#2E74B5")
    _center_text(draw, box, text, _font(17, True))


def _class_box(draw, box, name: str, attrs: list[str], ops: list[str] | None = None) -> None:
    ops = ops or []
    x1, y1, x2, y2 = box
    draw.rectangle(box, fill="#FFFFFF", outline="#2E74B5", width=2)
    draw.rectangle((x1, y1, x2, y1 + 40), fill="#EFF6FF", outline="#2E74B5", width=2)
    _center_text(draw, (x1, y1 + 2, x2, y1 + 38), name, _font(17, True))
    attr_y = y1 + 52
    for attr in attrs:
        draw.text((x1 + 14, attr_y), attr, font=_font(15), fill="#1F2933")
        attr_y += 22
    sep_y = y2 - max(34, 22 * len(ops) + 14)
    draw.line((x1, sep_y, x2, sep_y), fill="#2E74B5", width=2)
    op_y = sep_y + 10
    for op in ops:
        draw.text((x1 + 14, op_y), op, font=_font(14), fill="#1F2933")
        op_y += 20


def _association(draw, start, end, start_mult: str, end_mult: str, label: str = "") -> None:
    draw.line([start, end], fill="#34495E", width=2)
    dx = end[0] - start[0]
    dy = end[1] - start[1]
    if abs(dy) > abs(dx) and dy > 0:
        start_pos = (start[0] + 8, start[1] + 8)
        end_pos = (end[0] + 8, end[1] - 28)
        label_pos = (start[0] + 12, (start[1] + end[1]) // 2 - 8)
    elif abs(dy) > abs(dx) and dy < 0:
        start_pos = (start[0] + 8, start[1] - 28)
        end_pos = (end[0] + 8, end[1] + 8)
        label_pos = (start[0] + 12, (start[1] + end[1]) // 2 - 8)
    else:
        start_pos = (start[0] + 8, start[1] - 24)
        end_pos = (end[0] - 44, end[1] - 24)
        label_pos = ((start[0] + end[0]) // 2 - 34, start[1] + 8)
    draw.text(start_pos, start_mult, font=_font(14), fill="#34495E")
    draw.text(end_pos, end_mult, font=_font(14), fill="#34495E")
    if label:
        draw.text(label_pos, label, font=_font(14), fill="#34495E")


def _component_box(draw, box, name: str, stereotype: str = "<<component>>", fill="#EAF7EF") -> None:
    _box(draw, box, f"{stereotype}\n{name}", fill=fill, font=_font(16, True), radius=4)
    x1, y1, _, _ = box
    draw.rectangle((x1 - 18, y1 + 22, x1 + 8, y1 + 42), fill="#FFFFFF", outline="#2E74B5", width=2)
    draw.rectangle((x1 - 18, y1 + 54, x1 + 8, y1 + 74), fill="#FFFFFF", outline="#2E74B5", width=2)


def _node_box(draw, box, name: str, stereotype: str = "<<node>>", fill="#F8F9FA") -> None:
    x1, y1, x2, y2 = box
    offset = 22
    draw.polygon([(x1 + offset, y1), (x2, y1), (x2, y2 - offset), (x2 - offset, y2), (x1, y2), (x1, y1 + offset)], fill=fill, outline="#2E74B5")
    draw.line((x1, y1 + offset, x1 + offset, y1), fill="#2E74B5", width=2)
    draw.line((x2 - offset, y2, x2, y2 - offset), fill="#2E74B5", width=2)
    _center_text(draw, (x1 + 8, y1 + 8, x2 - 8, y2 - 8), f"{stereotype}\n{name}", _font(16, True))


def generate_uml_diagrams() -> dict[str, Path]:
    paths: dict[str, Path] = {}
    blue = "#2E74B5"
    pale = "#EFF6FF"
    green = "#EAF7EF"
    amber = "#FFF7E6"

    path, img, draw = _new_diagram("uml_use_case.png", (1700, 1050))
    _title(draw, "Use Case Diagram - VolunteerHub", 1700)
    draw.rectangle((330, 110, 1390, 930), outline=blue, width=3)
    draw.text((365, 125), "System boundary: VolunteerHub", font=_font(20, True), fill="#1F4E79")
    actors = [("Guest", 170, 160), ("Volunteer", 170, 445), ("Organizer", 170, 705), ("Admin", 1530, 250), ("Sponsor", 1530, 610)]
    for label, x, y in actors:
        _actor(draw, x, y, label)
    use_cases = [
        ("Browse events", (430, 170, 700, 250)), ("Verify certificate", (430, 295, 700, 375)),
        ("Register / login", (650, 430, 920, 510)), ("Manage profile / KYC", (650, 545, 920, 625)),
        ("Register event", (650, 660, 920, 740)), ("Check-in / checkout", (650, 775, 920, 855)),
        ("Create/manage event", (1010, 470, 1280, 550)), ("Manage shifts & attendance", (1010, 600, 1280, 680)),
        ("Campaign & donation", (1010, 730, 1300, 810)), ("Sponsorship proposal", (1010, 835, 1310, 915)),
        ("Review / monitor / export", (1010, 250, 1280, 330)),
    ]
    centers = {}
    use_case_boxes = {}
    for text, box in use_cases:
        _oval(draw, box, text, fill=pale)
        centers[text] = ((box[0] + box[2]) // 2, (box[1] + box[3]) // 2)
        use_case_boxes[text] = box
    links = [
        ((250, 230), "Browse events", "left"), ((250, 230), "Verify certificate", "left"),
        ((250, 515), "Register / login", "left"), ((250, 515), "Manage profile / KYC", "left"),
        ((250, 515), "Register event", "left"), ((250, 515), "Check-in / checkout", "left"), ((250, 515), "Campaign & donation", "left"),
        ((250, 775), "Create/manage event", "left"), ((250, 775), "Manage shifts & attendance", "left"), ((250, 775), "Campaign & donation", "left"), ((250, 775), "Sponsorship proposal", "left"),
        ((1450, 320), "Review / monitor / export", "right"),
        ((1450, 680), "Sponsorship proposal", "right"),
    ]
    for start, name, side in links:
        box = use_case_boxes[name]
        end = (box[0], (box[1] + box[3]) // 2) if side == "left" else (box[2], (box[1] + box[3]) // 2)
        draw.line([start, end], fill="#7895B2", width=2)
    include_rels = []
    for source, target in include_rels:
        sx, sy = centers[source]
        tx, ty = centers[target]
        _dashed_dependency(draw, (sx + 60, sy - 30), (tx - 40, ty + 30), "<<include>>")
    img.save(path); paths["use_case"] = path

    path, img, draw = _new_diagram("uml_activity_registration.png", (1500, 1200))
    _title(draw, "Activity Diagram - Đăng ký và điểm danh sự kiện", 1500)
    _initial_node(draw, 750, 115)
    actions = [
        ("Volunteer chọn event", (580, 150, 920, 210), pale),
        ("Gửi đăng ký", (580, 275, 920, 335), pale),
        ("Event Approved,\ncòn hạn, còn chỗ?", (560, 395, 940, 475), amber),
        ("KYC/skill/shift\nhợp lệ?", (560, 545, 940, 625), amber),
        ("Tạo Registration Pending", (560, 700, 940, 760), green),
        ("Organizer Confirm", (560, 830, 940, 890), green),
        ("QR/GPS/time window\nhợp lệ?", (560, 950, 940, 1030), amber),
        ("Ghi check-in/out,\ntính giờ", (560, 1060, 940, 1120), green),
    ]
    prev = (750, 128)
    for text, box, fill in actions:
        if "?" in text:
            _decision(draw, box, text, fill=fill)
        else:
            _box(draw, box, text, fill=fill, font=_font(18))
        center_top = ((box[0] + box[2]) // 2, box[1])
        center_bottom = ((box[0] + box[2]) // 2, box[3])
        _arrow(draw, prev, center_top)
        if "?" in text and "QR" not in text:
            draw.text((805, box[3] + 8), "[yes]", font=_font(14), fill="#34495E")
        prev = center_bottom
    _final_node(draw, 750, 1138)
    _arrow(draw, (750, 1120), (750, 1125))
    exceptions = [
        ((940, 435), (1010, 435), "Từ chối + thông báo lỗi", "[no]"),
        ((940, 585), (1010, 585), "Yêu cầu bổ sung KYC/skill", "[no]"),
        ((940, 990), (1010, 990), "Reject check-in", "[no]"),
    ]
    for start, end, text, guard in exceptions:
        draw.text((start[0] + 12, start[1] - 24), guard, font=_font(14), fill="#C0392B")
        _arrow(draw, start, end, fill="#C0392B")
        _box(draw, (1010, end[1] - 30, 1320, end[1] + 30), text, fill="#FDEDEC", outline="#C0392B", font=_font(17))
    img.save(path); paths["activity_registration"] = path

    path, img, draw = _new_diagram("uml_component_deployment.png", (1600, 950))
    _title(draw, "Component / Deployment Diagram - VolunteerHub", 1600)
    _node_box(draw, (60, 150, 380, 330), "Client Browser\nReact/Vite SPA\nport 3000", "<<device>>", pale)
    _node_box(draw, (470, 125, 800, 350), "API Gateway Host\nOcelot port 5000", "<<executionEnvironment>>", amber)
    service_node = (900, 70, 1280, 825)
    _node_box(draw, service_node, "Backend Services Host", "<<executionEnvironment>>", "#F8FBFF")
    _component_box(draw, (945, 125, 1235, 235), "AuthService\nport 5002", fill=green)
    _component_box(draw, (945, 300, 1235, 410), "EventService\nport 5003", fill=green)
    _component_box(draw, (945, 475, 1235, 585), "FinanceService\nport 5004", fill=green)
    _component_box(draw, (945, 650, 1235, 760), "Legacy API\nport 5001", fill="#F3F4F6")
    _component_box(draw, (500, 500, 770, 600), "SignalR Hub\n/hubs/channel", fill=pale)
    _node_box(draw, (1340, 330, 1550, 560), "SQL Server\nVolunteerHub DB", "<<database>>", "#FFFFFF")
    for y in [240, 160, 350, 540, 730]:
        if y == 240:
            _arrow(draw, (380, 240), (470, 240))
        elif y == 160:
            _arrow(draw, (800, 240), (945, 180))
        elif y == 350:
            _arrow(draw, (800, 240), (945, 355))
        elif y == 540:
            _arrow(draw, (800, 240), (945, 530))
        else:
            _arrow(draw, (800, 240), (945, 705))
    for y in [160, 350, 540, 730]:
        _arrow(draw, (1235, y), (1340, 445))
    _arrow(draw, (635, 500), (1090, 410), fill="#6C5CE7")
    img.save(path); paths["component"] = path

    path, img, draw = _new_diagram("uml_class_domain.png", (1700, 1100))
    _title(draw, "Class / Domain Model - VolunteerHub", 1700)
    classes = [
        ("User", ["- Id: int", "- UserName: string", "- UserType: enum", "- IsActive: bool"], ["+ Login()"], (80, 140, 350, 320)),
        ("VolunteerProfile", ["- KycStatus: string", "- TotalHours: decimal"], ["+ SubmitKyc()"], (80, 430, 350, 590)),
        ("Event", ["- Status: string", "- StartDate: DateTime", "- EndDate: DateTime", "- QrCode: string"], ["+ Approve()", "+ Complete()"], (540, 140, 850, 340)),
        ("Registration", ["- Status: string", "- AttendedAt: DateTime?", "- VolunteerHours: decimal"], ["+ CheckIn()", "+ CheckOut()"], (540, 430, 850, 630)),
        ("WorkShift", ["- StartTime: DateTime", "- EndTime: DateTime", "- Capacity: int"], [], (940, 430, 1230, 590)),
        ("Certificate", ["- CertificateCode: string", "- VolunteerHours: decimal"], ["+ Verify()"], (540, 780, 850, 950)),
        ("SupportCampaign", ["- Status: string", "- TargetAmount: decimal"], ["+ Close()", "+ Report()"], (1260, 140, 1570, 320)),
        ("Donation", ["- Status: string", "- Amount: decimal"], ["+ Confirm()"], (1320, 430, 1620, 590)),
        ("SponsorshipProposal", ["- Type: string", "- Status: string", "- ActualAmount: decimal"], ["+ MarkReceived()"], (1190, 780, 1600, 960)),
    ]
    for name, attrs, ops, box in classes:
        _class_box(draw, box, name, attrs, ops)
    _association(draw, (350, 230), (540, 230), "1", "0..*", "organizes")
    _association(draw, (215, 320), (215, 430), "1", "0..1", "profile")
    _association(draw, (350, 510), (540, 530), "1", "0..*", "registers")
    _association(draw, (695, 340), (695, 430), "1", "0..*", "has")
    _association(draw, (850, 530), (940, 530), "0..*", "0..1", "shift")
    _association(draw, (695, 630), (695, 780), "1", "0..1", "issues")
    _association(draw, (850, 230), (1260, 230), "1", "0..*", "campaign")
    _association(draw, (1415, 320), (1415, 430), "1", "0..*", "donations")
    _association(draw, (850, 865), (1190, 865), "0..*", "0..*", "sponsor impact")
    img.save(path); paths["class_domain"] = path

    path, img, draw = _new_diagram("uml_sequence_approve_event.png", (1650, 1050))
    _title(draw, "Sequence Diagram - Admin duyệt sự kiện", 1650)
    lifelines = [("Admin UI", 150), ("Gateway", 430), ("EventController", 710), ("EventService", 990), ("DbContext", 1250), ("Notification", 1480)]
    for name, x in lifelines:
        _box(draw, (x - 90, 120, x + 90, 175), name, fill=pale, font=_font(16, True))
        draw.line((x, 175, x, 960), fill="#BDC3C7", width=2)
    for x, y1, y2 in [(150, 230, 895), (430, 240, 880), (710, 320, 815), (990, 400, 800), (1250, 480, 660), (1480, 720, 765)]:
        _activation(draw, x, y1, y2)
    messages = [
        (150, 430, 240, "PUT /events/{id}/approve"),
        (430, 710, 320, "route request"),
        (710, 990, 400, "ApproveAsync(id, adminId)"),
        (990, 1250, 480, "Load Event + Organizer"),
        (990, 1250, 640, "Set Approved, QR, Channel"),
        (990, 1480, 720, "Send notification"),
    ]
    for x1, x2, y, label in messages:
        _arrow(draw, (x1, y), (x2, y), width=2)
        draw.text((min(x1, x2) + 12, y - 28), label, font=_font(15), fill="#2C3E50")
    _return_arrow(draw, (1250, 560), (990, 560), "Event aggregate")
    _return_arrow(draw, (990, 800), (710, 800), "Result DTO")
    _return_arrow(draw, (710, 880), (150, 880), "200 OK + toast")
    img.save(path); paths["sequence_approve"] = path

    path, img, draw = _new_diagram("uml_state_machines.png", (1650, 1100))
    _title(draw, "State Machine Diagram - Event / Registration / Finance", 1650)
    rows = [
        ("Event", ["Pending", "Approved", "Completed"], [("Rejected", 0), ("Cancelled", 1)]),
        ("Registration", ["Pending", "Confirmed", "CheckedOut"], [("CancelRequested", 1), ("Cancelled", 0)]),
        ("Donation", ["PendingConfirmation", "Confirmed"], [("Rejected", 0), ("Cancelled", 0)]),
        ("Proposal", ["Pending", "Accepted", "Received", "Reported"], [("Rejected", 0), ("Cancelled", 1)]),
    ]
    y = 140
    for title, main, alternatives in rows:
        draw.text((70, y + 25), title, font=_font(21, True), fill="#1F4E79")
        _initial_node(draw, 250, y + 35)
        x = 300
        prev = None
        main_boxes = []
        for state in main:
            box = (x, y, x + 210, y + 70)
            _box(draw, box, state, fill=green, font=_font(17, True))
            main_boxes.append(box)
            if prev:
                _arrow(draw, (prev[2], y + 35), (box[0], y + 35), width=2)
            prev = box
            x += 270
        _arrow(draw, (263, y + 35), (300, y + 35), width=2)
        _final_node(draw, prev[2] + 55, y + 35)
        _arrow(draw, (prev[2], y + 35), (prev[2] + 38, y + 35), width=2)
        alt_y = y + 98
        alt_x = 360
        previous_alt = None
        for alt, source_idx in alternatives:
            box = (alt_x, alt_y, alt_x + 230, alt_y + 64)
            _box(draw, box, fill="#FDEDEC", outline="#C0392B", text=alt, font=_font(16, True))
            if previous_alt is None:
                src = main_boxes[min(source_idx, len(main_boxes) - 1)]
                _arrow(draw, ((src[0] + src[2]) // 2, src[3]), ((box[0] + box[2]) // 2, box[1]), fill="#C0392B", width=2)
            if previous_alt:
                _arrow(draw, (previous_alt[2], alt_y + 32), (box[0], alt_y + 32), fill="#C0392B", width=2)
            previous_alt = box
            alt_x += 300
        y += 245
    img.save(path); paths["state"] = path

    return paths


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.35) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = MUTED


def add_page_section(doc: Document, heading: str, level: int = 2) -> None:
    doc.add_heading(heading, level=level)


def intro_common(doc: Document, kind: str) -> None:
    doc.add_heading("GIỚI THIỆU", level=1)
    doc.add_heading("Mục đích", level=2)
    doc.add_paragraph(
        f"Tài liệu {kind} mô tả có hệ thống phạm vi, yêu cầu và cơ sở thiết kế của phần mềm VolunteerHub. "
        "Nội dung được viết như một báo cáo đồ án hoàn chỉnh: trình bày bài toán, cách nhóm hiểu nghiệp vụ, "
        "phạm vi xây dựng, các chức năng chính và cách hệ thống được tổ chức để triển khai."
    )
    doc.add_heading("Phạm vi", level=2)
    doc.add_paragraph(
        "Phạm vi tài liệu bao gồm hệ thống web quản lý hoạt động tình nguyện: người dùng công khai, tình nguyện viên, "
        "nhà tổ chức, nhà tài trợ và quản trị viên. Tài liệu không mô tả chi tiết các module legacy bán hàng cũ ngoài vai trò fallback."
    )
    doc.add_heading("Cách trình bày tài liệu", level=2)
    doc.add_paragraph(
        "Các phần trong báo cáo được sắp xếp theo luồng tự nhiên của một đồ án phần mềm. Trước hết tài liệu nêu bối cảnh "
        "và mục tiêu của hệ thống, sau đó đi vào người dùng, nghiệp vụ, chức năng, dữ liệu, giao diện và tiêu chí nghiệm thu. "
        "Những chi tiết kỹ thuật được đưa vào đúng vị trí cần thiết để người đọc hiểu được cách hệ thống vận hành mà không phải "
        "đọc từng dòng mã nguồn."
    )
    doc.add_heading("Từ và thuật ngữ", level=2)
    rows = [
        ("Volunteer", "Tình nguyện viên đăng ký, tham gia và nhận ghi nhận đóng góp."),
        ("Organizer", "Nhà tổ chức tạo sự kiện, quản lý đăng ký, điểm danh, hoàn thành sự kiện."),
        ("Sponsor", "Nhà tài trợ gửi đề nghị tài trợ và theo dõi hỗ trợ cho sự kiện."),
        ("Admin", "Quản trị viên duyệt dữ liệu, vận hành hệ thống và giám sát."),
        ("KYC", "Quy trình xác minh danh tính tình nguyện viên bằng minh chứng."),
        ("Campaign", "Đợt kêu gọi ủng hộ cá nhân gắn với sự kiện."),
        ("Sponsorship Proposal", "Đề nghị tài trợ doanh nghiệp giữa sponsor và organizer."),
        ("Volunteer Passport", "Hồ sơ tổng hợp sự kiện đã tham gia, giờ tình nguyện, chứng chỉ và huy hiệu."),
    ]
    add_table(doc, ["Thuật ngữ", "Diễn giải"], rows, [2200, 7160])


def add_natural_project_context(doc: Document, perspective: str) -> None:
    doc.add_heading("BỐI CẢNH VÀ ĐỊNH HƯỚNG XÂY DỰNG", level=1)
    doc.add_heading("Lý do hình thành hệ thống", level=2)
    doc.add_paragraph(
        "Các hoạt động tình nguyện thường bắt đầu từ những lời kêu gọi trên mạng xã hội hoặc các nhóm cộng đồng nhỏ. "
        "Cách tổ chức này nhanh và gần gũi, nhưng khi số lượng sự kiện tăng lên thì bắt đầu xuất hiện nhiều vấn đề: người tham gia "
        "khó biết tổ chức nào đáng tin cậy, nhà tổ chức khó tìm đúng người có kỹ năng phù hợp, nhà tài trợ khó theo dõi đóng góp "
        "của mình, còn tình nguyện viên lại thiếu một hồ sơ chính thống để ghi nhận quá trình cống hiến."
    )
    doc.add_paragraph(
        "VolunteerHub được xây dựng để gom các hoạt động đó về một cổng thống nhất. Hệ thống không chỉ đăng tin sự kiện, "
        "mà còn hỗ trợ tuyển chọn, xác minh, điểm danh, ghi nhận giờ tình nguyện, cấp chứng chỉ, tiếp nhận tài trợ và công khai "
        "kết quả đóng góp. Mục tiêu của đồ án là tạo ra một nền tảng đủ rõ ràng để vận hành trong phạm vi học thuật, đồng thời "
        "có cấu trúc đủ tốt để mở rộng nếu triển khai thực tế."
    )

    doc.add_heading("Cách nhóm hiểu bài toán", level=2)
    doc.add_paragraph(
        "Bài toán không chỉ là quản lý danh sách sự kiện. Trọng tâm của VolunteerHub nằm ở niềm tin và điều phối. "
        "Niềm tin được tạo ra bằng hồ sơ tổ chức, cơ chế duyệt, đánh giá hai chiều, lịch sử tham gia, chứng chỉ có mã xác thực "
        "và báo cáo tác động. Điều phối được giải quyết bằng thông tin kỹ năng, địa điểm, ca làm việc, quy trình đăng ký, phỏng vấn "
        "khi cần và điểm danh bằng QR hoặc GPS."
    )
    doc.add_paragraph(
        "Từ cách nhìn đó, hệ thống được chia thành các nhóm chức năng xoay quanh từng vai trò: tình nguyện viên cần tìm việc phù hợp "
        "và được ghi nhận; nhà tổ chức cần tuyển người, vận hành sự kiện và báo cáo kết quả; nhà tài trợ cần theo dõi dòng hỗ trợ; "
        "quản trị viên cần kiểm duyệt, xử lý dữ liệu nhạy cảm và giám sát hoạt động chung."
    )

    doc.add_heading("Định hướng giải pháp", level=2)
    add_table(
        doc,
        ["Mảng nghiệp vụ", "Cách hệ thống xử lý"],
        [
            (
                "Tìm kiếm và đăng ký sự kiện",
                "Sự kiện có thông tin kỹ năng, thời gian, địa điểm, số lượng và trạng thái duyệt. Tình nguyện viên xem danh sách, lọc theo nhu cầu, mở bản đồ và đăng ký khi đủ điều kiện.",
            ),
            (
                "Xác minh và uy tín",
                "Hồ sơ tổ chức, KYC tình nguyện viên, xác minh kỹ năng và đánh giá sau sự kiện giúp hệ thống hạn chế gian lận, đồng thời tạo dữ liệu tin cậy cho các lần tham gia sau.",
            ),
            (
                "Vận hành tại hiện trường",
                "Nhà tổ chức quản lý danh sách đăng ký, ca làm việc, điểm danh, check-out và walk-in. QR/GPS giúp xác nhận sự có mặt nhưng vẫn có đường xử lý thủ công khi tình huống thực tế phát sinh.",
            ),
            (
                "Ghi nhận đóng góp",
                "Sau khi sự kiện hoàn thành, hệ thống tổng hợp giờ tham gia, cấp chứng chỉ, trao huy hiệu và cập nhật Volunteer Passport để tình nguyện viên có hồ sơ đóng góp rõ ràng.",
            ),
            (
                "Tài trợ và minh bạch",
                "Campaign, donation và sponsorship proposal được quản lý theo trạng thái. Chỉ các khoản đã xác nhận mới được tính vào kết quả công khai, giúp người xem phân biệt cam kết và đóng góp thực nhận.",
            ),
            (
                "Quản trị và giám sát",
                "Admin duyệt tổ chức, duyệt KYC/kỹ năng, quản lý user, rating, export và audit log. Những thao tác nhạy cảm được đưa qua backend thay vì chỉ dựa vào giao diện.",
            ),
        ],
        [2600, 6760],
    )

    doc.add_heading("Hiện trạng triển khai trong đồ án", level=2)
    doc.add_paragraph(
        "Trong phạm vi đồ án, hệ thống đã có frontend React/Vite, API gateway, các service backend .NET, cơ sở dữ liệu SQL Server "
        "và một ứng dụng mobile phụ trợ. Các controller nghiệp vụ được tổ chức theo miền chức năng; một số controller gốc được đặt "
        "trong project APIService và được liên kết sang các service tương ứng khi chạy theo mô hình gateway. Cách tổ chức này giúp "
        "nhóm vừa giữ được cấu trúc triển khai theo service, vừa tránh nhân đôi code trong giai đoạn phát triển."
    )
    doc.add_paragraph(
        "Báo cáo " + perspective + " vì vậy được viết theo trạng thái thực tế của đồ án: phần nào đã có luồng nghiệp vụ và màn hình "
        "thì mô tả như chức năng hiện hữu; phần nào phụ thuộc môi trường thật như GPS/camera, tải lớn hoặc chia sẻ qua nền tảng ngoài "
        "thì được trình bày như điểm cần kiểm thử thêm khi nghiệm thu."
    )


def add_source_requirement_alignment(doc: Document, perspective: str) -> None:
    doc.add_heading("ĐỐI CHIẾU YÊU CẦU GỐC WEBSITE VOLUNTEER HUB", level=1)
    doc.add_paragraph(
        "Phần này đối chiếu trực tiếp với file Tài liệu CNPM/Website Volunteer Hub.docx. "
        "Mục tiêu là bảo đảm tài liệu không chỉ đúng template mà còn phủ đủ bối cảnh, actor, chức năng và NFR đã được giao."
    )
    rows = [
        (
            "Bối cảnh NGO/CLB và cá nhân thiện nguyện",
            "Hệ sinh thái minh bạch kết nối NGO, câu lạc bộ, organizer, volunteer và sponsor.",
            "Phần tổng quan hệ thống, actor, module Organizer Verification và Finance.",
        ),
        (
            "Tối ưu điều phối theo kỹ năng/vị trí",
            "Tìm kiếm/lọc sự kiện theo kỹ năng, vị trí; gợi ý sự kiện theo kỹ năng và bản đồ.",
            "EventList/EventDetail, MapView/Leaflet, registration rule về skill/shift/KYC.",
        ),
        (
            "Rating hai chiều",
            "Volunteer đánh giá organizer và organizer đánh giá volunteer sau sự kiện.",
            "Rating module, Admin moderation, API /ratings và tiêu chí nghiệm thu rating.",
        ),
        (
            "Số hóa đóng góp",
            "Volunteer Passport, tổng giờ tình nguyện, lịch sử dự án, certificate và badge.",
            "VolunteerProfile.TotalVolunteerHours, Registrations.VolunteerHours, CertificateService, BadgeService.",
        ),
        (
            "Hồ sơ kỹ năng/nhóm máu/ngôn ngữ/sở thích",
            "Profile volunteer lưu kỹ năng, bio, avatar, nhóm máu, ngôn ngữ, sở thích và minh chứng kỹ năng.",
            "ProfileService, VolunteerProfiles, VolunteerSkills, UI Profile.",
        ),
        (
            "Đăng tin tuyển dụng sự kiện",
            "Organizer mô tả công việc, kỹ năng, thời gian, địa điểm, số lượng cần tuyển.",
            "Events entity, EventForm, WorkShift và RequiredSkillIds.",
        ),
        (
            "Phê duyệt volunteer và phỏng vấn trực tuyến",
            "Organizer xem hồ sơ, hẹn/đổi/hủy/chấm kết quả phỏng vấn trước khi confirm nếu event yêu cầu.",
            "RequiresInterview, InterviewService, interview endpoints, ManageEvent registration tab.",
        ),
        (
            "Bản đồ tình nguyện",
            "Hiển thị sự kiện quanh vị trí người dùng và liên kết bản đồ/địa điểm.",
            "EventList map view, EventDetail location map, latitude/longitude.",
        ),
        (
            "Điểm danh QR/GPS và quản lý ca",
            "Check-in bằng QR hoặc định vị GPS, checkout tính giờ; phân bổ theo ca/vị trí công việc.",
            "RegistrationService check-in/out, WorkShifts, QR rotation, GPS radius.",
        ),
        (
            "Chứng chỉ PDF có mã QR định danh",
            "Sau khi sự kiện hoàn thành, tự động tạo chứng chỉ PDF/ảnh có QR verify code.",
            "CertificateService, CertificateCode unique, public verify endpoint.",
        ),
        (
            "Sponsor đóng góp tài chính/nhu yếu phẩm",
            "Sponsor theo dõi dự án và đóng góp nguồn lực bằng tiền hoặc hiện vật/nhu yếu phẩm.",
            "SponsorshipProposal.Type, InKindDescription/ActualReceivedAmount, campaign/donation/proposal impact.",
        ),
        (
            "Admin kiểm duyệt pháp lý, khiếu nại và báo cáo tác động",
            "Admin duyệt tổ chức, vận hành rating/moderation, giám sát finance và audit/export báo cáo.",
            "OrganizerVerification, Admin monitoring, ratings hide/delete, AuditLogs, export.",
        ),
        (
            "NFR: minh bạch, bảo mật, hiệu năng, mobile, chia sẻ mạng xã hội",
            "Public impact, JWT/role, pagination/rate limit/export limit, responsive/mobile, share event/certificate/badge.",
            "NFR section, Gateway/Auth middleware, frontend responsive screens, share links trong public/achievement flows.",
        ),
        (
            "Chiến dịch khẩn cấp/cứu trợ thiên tai",
            "Hệ thống phải chịu được lượng truy cập lớn khi có chiến dịch khẩn cấp.",
            "NFR performance: pagination, caching/rate limit định hướng, export limit, deployment scaling note.",
        ),
    ]
    add_table(doc, ["Yêu cầu gốc", "Diễn giải trong " + perspective, "Thiết kế/hiện thực liên quan"], rows, [2600, 3400, 3360])


def add_nfr_design_response(doc: Document) -> None:
    doc.add_heading("THIẾT KẾ CHẤT LƯỢNG HỆ THỐNG", level=1)
    doc.add_paragraph(
        "Ngoài các chức năng nghiệp vụ, VolunteerHub cần tạo cảm giác tin cậy khi được sử dụng trong cộng đồng. "
        "Thiết kế vì vậy chú ý đến tính minh bạch, bảo mật thông tin cá nhân, khả năng phục vụ khi có nhiều người truy cập, "
        "trải nghiệm trên thiết bị di động và khả năng chia sẻ các nội dung công khai."
    )
    add_table(
        doc,
        ["Khía cạnh chất lượng", "Cách thiết kế trong VolunteerHub", "Điểm cần lưu ý khi vận hành"],
        [
            (
                "Minh bạch",
                "Thông tin tổ chức, sự kiện, kết quả đóng góp và mã chứng chỉ được đưa ra các trang công khai ở mức phù hợp. Các khoản tài trợ chỉ được tính vào impact khi đã được xác nhận.",
                "Cần giữ nguyên nguyên tắc không cộng tiền/công sức ở trạng thái chờ xác nhận; thao tác duyệt và điều chỉnh nên có audit để truy lại khi cần.",
            ),
            (
                "Bảo mật",
                "Backend kiểm đăng nhập, vai trò, trạng thái tài khoản và quyền sở hữu dữ liệu. Thông tin nhạy cảm như password, salt, hồ sơ KYC và minh chứng không được trả tự do qua API công khai.",
                "Frontend chỉ hỗ trợ điều hướng, không được xem là lớp bảo vệ chính. Các API admin, finance, KYC và upload cần tiếp tục được kiểm quyền ở backend.",
            ),
            (
                "Hiệu năng",
                "Các danh sách chính dùng phân trang và lọc; export có giới hạn; API nhạy cảm có rate limit. Gateway giúp tách luồng truy cập theo nhóm chức năng khi cần mở rộng.",
                "Nếu triển khai cho chiến dịch khẩn cấp quy mô lớn, public event list và impact dashboard nên bổ sung caching/monitoring thực tế.",
            ),
            (
                "Khả dụng trên mobile",
                "Các luồng hay dùng tại hiện trường như xem sự kiện, mở bản đồ, đăng ký, QR/GPS check-in, check-out và xem chứng chỉ được thiết kế để thao tác được trên màn hình nhỏ.",
                "GPS, camera và mạng di động phụ thuộc thiết bị thật; khi nghiệm thu cần chạy thêm trên điện thoại ngoài môi trường phòng máy.",
            ),
            (
                "Kết nối và chia sẻ",
                "Sự kiện, hồ sơ công khai và chứng chỉ có đường dẫn có thể chia sẻ. Volunteer có thể chia sẻ thành tích hoặc link verify mà người nhận không cần đăng nhập để xem phần công khai.",
                "Phần chia sẻ hiện tập trung vào URL công khai và nút chia sẻ phía client; chưa xem là tích hợp sâu với API của mạng xã hội.",
            ),
        ],
        [2100, 4400, 2860],
    )

    doc.add_heading("Tác động đến kiến trúc", level=2)
    doc.add_paragraph(
        "Các yêu cầu chất lượng trên ảnh hưởng trực tiếp đến cách tổ chức backend và frontend. Những luồng cần bảo mật được đặt sau "
        "JWT, role policy và middleware kiểm tra trạng thái tài khoản. Những luồng có nguy cơ tải lớn như danh sách sự kiện, dashboard "
        "và export được thiết kế theo hướng phân trang, giới hạn dữ liệu và có thể mở rộng bằng cache. Các thao tác tại hiện trường được "
        "tách thành màn hình rõ ràng để organizer và volunteer xử lý nhanh trong lúc sự kiện đang diễn ra."
    )
    add_table(
        doc,
        ["Khu vực hệ thống", "Vai trò trong thiết kế chất lượng", "Ghi chú triển khai"],
        [
            ("Gateway và backend services", "Định tuyến API, kiểm JWT, rate limit và tách luồng theo miền nghiệp vụ.", "Các service dùng chung entity/service layer trong phạm vi đồ án để giảm trùng lặp."),
            ("Event và Registration", "Giữ tính đúng đắn của đăng ký, ca làm, check-in/check-out và hoàn thành sự kiện.", "QR/GPS là lớp xác nhận nhanh; organizer vẫn có thao tác thủ công có kiểm soát."),
            ("Finance", "Giữ minh bạch cho donation, campaign và sponsorship proposal.", "Chỉ trạng thái đã xác nhận/đã nhận mới được tính vào impact công khai."),
            ("Certificate và Badge", "Ghi nhận đóng góp bằng mã xác thực, chứng chỉ và huy hiệu.", "Trang verify public chỉ nên hiển thị dữ liệu cần thiết."),
            ("Frontend web/mobile", "Tối ưu thao tác thực tế cho volunteer, organizer, sponsor và admin.", "Các lỗi GPS, camera, mạng và quyền truy cập cần thông báo rõ cho người dùng."),
        ],
        [2500, 4300, 2560],
    )


def add_source_code_round9_audit(doc: Document, perspective: str) -> None:
    doc.add_heading("ĐỐI CHIẾU VÒNG 9 VỚI SOURCE CODE", level=1)
    doc.add_paragraph(
        f"Phần này rà soát {perspective} theo source code hiện tại, tập trung vào các khẳng định dễ bị lệch: "
        "kiến trúc service/gateway, route API, service layer, entity/database, frontend/mobile và test evidence. "
        "Kết quả được ghi theo hướng thận trọng: cái đã có trong source được nêu như hiện trạng; cái là định hướng mở rộng "
        "được ghi là extension hoặc rủi ro kiểm soát."
    )
    add_table(
        doc,
        ["Hạng mục kiểm tra", "Dấu vết trong source code", "Kết luận đưa vào tài liệu"],
        [
            (
                "API Gateway và service split",
                "BaseCore.ApiGateway/ocelot.json định tuyến /api/auth,/api/profile,/api/users,/api/skills sang 5002; /api/events,/api/certificates,/api/channels,/api/dashboard sang 5003; finance sang 5004; fallback /api/* sang 5001.",
                "Ghi là kiến trúc gateway/service-ready; source controller gốc đặt trong APIService và được link sang service qua .csproj.",
            ),
            (
                "Controller/API hiện hữu",
                "BaseCore.APIService/Controllers có Events, Registrations, Certificates, WorkShifts, Profile, OrganizerVerification, Admin, Ratings, Channels, SupportCampaign, SponsorshipProposal, Uploads.",
                "Bảng API khớp các route thật; route đặc thù như /api/events/{id}/qr/rotate, /api/events/{eventId}/self-checkin, /api/profile/passport được giữ.",
            ),
            (
                "Business service layer",
                "BaseCore.Services/VolunteerHub gồm EventService, RegistrationService, InterviewService, CertificateService, BadgeService, ChannelService, NotificationService, AuditLogService.",
                "Thiết kế controller mỏng/service chứa nghiệp vụ có cơ sở source; sequence và class diagram bám vào các service này.",
            ),
            (
                "Entity/database",
                "BaseCore.Entities và MySqlDbContext có User, VolunteerProfile, OrganizerVerification, Event, WorkShift, Registration, Certificate, Badge, SupportCampaign, IndividualDonation, SponsorProfile, SponsorshipProposal, Rating, Notification, AuditLog.",
                "Domain model và state machine có bám entity thật; các ràng buộc unique/index/relationship được mô tả ở mức thiết kế, không thay thế migration chi tiết.",
            ),
            (
                "Bảo mật/NFR",
                "Program.cs ở API/Event/Finance dùng JWT Bearer, IsActive middleware và AddRateLimiter; TokenHelper hash password bằng salt; controller dùng Authorize/role.",
                "Security/performance NFR được coi là đã có cơ chế thiết kế và implementation nền tảng.",
            ),
            (
                "QR/GPS/check-in",
                "RegistrationService validate QR/GPS/time window; Event.CheckInRadiusKm, Latitude, Longitude; WebClient có VolunteerCheckInModal, CheckInTab, html5-qrcode, QRCodeCanvas.",
                "Yêu cầu điểm danh thông minh được ghi là có implementation web; kiểm GPS/camera thật vẫn cần UAT trên thiết bị thật.",
            ),
            (
                "Certificate/badge/passport",
                "CertificateService, CertificatesController, /api/certificates/{code}, /api/profile/passport, Passport/Achievements/MyCertificates/VerifyCertificate UI.",
                "Yêu cầu số hóa đóng góp có trace đủ từ backend đến UI.",
            ),
            (
                "Mobile",
                "BaseCore.MobileApp có Expo routes events, checkin, profile, activity, my-sponsorships, admin và interview.",
                "Mobile được nêu như app/front-end phụ hỗ trợ luồng chính; không thay cho kiểm thử thiết bị thật.",
            ),
            (
                "Kết nối/social sharing",
                "EventDetail/Profile dùng navigator.share/copy link; MyCertificates/MyBadges tạo LinkedIn/Facebook share URLs.",
                "Chỉ ghi là share link/social entry. Không ghi là tích hợp social API server-side đầy đủ.",
            ),
            (
                "Test evidence",
                "BaseCore.WebClient/tests/e2e và TestResults/run-20260519-business ghi business/API probe; manual-critical có các case PARTIAL/SKIP cho GPS/camera/thiết bị thật.",
                "Checklist kiểm thử trong tài liệu bám test hiện có và giữ rõ phần cần UAT thêm.",
            ),
        ],
        [2100, 4300, 2960],
    )


def add_se_methodology_for_srs(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("CƠ SỞ PHƯƠNG PHÁP ĐẶC TẢ YÊU CẦU", level=1)
    doc.add_paragraph(
        "Phần này được xây dựng sau khi đối chiếu hai tài liệu nền tảng trong thư mục Tài liệu CNPM: "
        "Essentials of Software Engineering, Third Edition và Software Engineering with UML. "
        "Hai tài liệu này cho thấy SRS không nên chỉ là danh sách chức năng, mà cần thể hiện quá trình "
        "hình thành yêu cầu, phân tích actor/use case, phân loại yêu cầu, kiểm chứng và truy vết."
    )
    doc.add_heading("Quy trình requirements engineering áp dụng", level=2)
    add_table(
        doc,
        ["Bước", "Ý nghĩa trong kỹ nghệ phần mềm", "Áp dụng cho VolunteerHub"],
        [
            ("Elicitation", "Thu thập nhu cầu từ stakeholder, tài liệu hiện trạng, hệ thống cũ và bối cảnh nghiệp vụ.",
             "Dùng Website Volunteer Hub.docx, docs, Context, source code, test scenarios và báo cáo mẫu CNPM làm nguồn."),
            ("Analysis", "Làm rõ mâu thuẫn, phân rã yêu cầu, xác định actor, use case, ràng buộc và ưu tiên.",
             "Tách Guest, Volunteer, Organizer, Sponsor, Admin; gom yêu cầu thành event, registration, finance, engagement, admin."),
            ("Specification", "Viết yêu cầu ở dạng có thể đọc, kiểm thử và truy vết.",
             "Mỗi FR có actor, input, output, rule, acceptance criteria; mỗi UC có main/alternative flow."),
            ("Validation", "Kiểm tra yêu cầu có đúng, đầy đủ, nhất quán và khả thi không.",
             "Đối chiếu với source controller/service, e2e tests, docs/kich-ban-nghiep-vu-thuc-te.md."),
            ("Management", "Theo dõi thay đổi yêu cầu, trạng thái và quan hệ với thiết kế/test.",
             "Dùng mã FR/UC, traceability matrix và mapping sang API/UI/entity."),
        ],
        [1450, 3550, 4360],
    )
    doc.add_heading("Phân loại yêu cầu", level=2)
    add_table(
        doc,
        ["Loại yêu cầu", "Cách hiểu", "Ví dụ trong VolunteerHub"],
        [
            ("Business requirement", "Mục tiêu nghiệp vụ hoặc giá trị hệ thống phải tạo ra.",
             "Minh bạch hóa hoạt động tình nguyện, ghi nhận đóng góp chính thống, kết nối đúng người đúng việc."),
            ("Functional requirement", "Hành vi hệ thống phải thực hiện để actor đạt mục tiêu.",
             "Volunteer đăng ký event, organizer xác nhận, admin duyệt KYC, sponsor gửi proposal."),
            ("Domain rule", "Quy tắc nghiệp vụ đặc thù miền bài toán.",
             "Chỉ proposal Received mới tính tài trợ; check-in không cộng giờ ngay; event có ca thì phải chọn ca."),
            ("Constraint", "Giới hạn công nghệ, môi trường hoặc vận hành.",
             ".NET 8, React/Vite, SQL Server, gateway Ocelot, JWT, dùng chung database trong phạm vi đồ án."),
            ("Nonfunctional requirement", "Yêu cầu vận hành, chất lượng, bảo mật, hiệu năng, khả dụng.",
             "Role-based access, IsActive middleware, audit log, responsive UI, pagination/export limit."),
        ],
        [2100, 3200, 4060],
    )
    doc.add_heading("Nguyên tắc actor và use case", level=2)
    doc.add_paragraph(
        "Theo cách tiếp cận UML, actor là vai trò nằm ngoài ranh giới hệ thống và tương tác với hệ thống để đạt mục tiêu. "
        "Use case mô tả chuỗi tương tác tạo ra kết quả có giá trị cho actor. Vì vậy tài liệu này không xem User, Event hay "
        "Registration là actor; đó là entity/class bên trong hệ thống. Các actor của VolunteerHub được đặt theo vai trò sử dụng: "
        "Guest, Volunteer, Organizer, Sponsor, Admin, cùng actor gián tiếp như Time/System Scheduler khi nói về auto-complete."
    )
    add_table(
        doc,
        ["Actor", "Kiểu actor", "Lý do đưa vào SRS"],
        [
            ("Guest", "Primary human actor", "Nhận giá trị từ việc xem sự kiện public và verify certificate."),
            ("Volunteer", "Primary human actor", "Tạo phần lớn flow nghiệp vụ: hồ sơ, đăng ký, điểm danh, donation, passport."),
            ("Organizer", "Primary human actor", "Vận hành event, registration, shift, campaign và proposal."),
            ("Sponsor", "Primary human actor", "Tài trợ doanh nghiệp và theo dõi tác động."),
            ("Admin", "Primary/secondary human actor", "Vận hành, kiểm duyệt, giám sát và xử lý ngoại lệ."),
            ("Time/System Scheduler", "Nonhuman actor", "Kích hoạt auto-complete event, auto-close campaign, reminder."),
        ],
        [1800, 2600, 4960],
    )
    doc.add_heading("Yêu cầu phi chức năng theo operational perspective", level=2)
    doc.add_paragraph(
        "Software Engineering with UML nhấn mạnh NFR thường khó nhìn thấy trong use case flow vì chúng không phải hành vi "
        "nghiệp vụ trực tiếp. Vì vậy tài liệu này tách NFR thành phần riêng, đồng thời gắn chúng bằng note/constraint vào "
        "use case và kiến trúc. Với VolunteerHub, các NFR quan trọng nhất là bảo mật, minh bạch, khả dụng mobile, hiệu năng "
        "danh sách lớn, auditability và maintainability."
    )


def add_se_methodology_for_design(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("CƠ SỞ PHƯƠNG PHÁP THIẾT KẾ", level=1)
    doc.add_paragraph(
        "Thiết kế chi tiết được tổ chức theo hướng trace từ problem space sang solution space và architectural space. "
        "Essentials of Software Engineering phân biệt architectural design, high-level design và detailed design; "
        "Software Engineering with UML bổ sung cách chuyển use case/activity sang class, sequence, state, persistence, "
        "component/deployment và test design. Với VolunteerHub, điều này dẫn đến cấu trúc tài liệu gồm kiến trúc service, "
        "module, class/service, entity/database, API, UI, state machine, sequence và triển khai."
    )
    doc.add_heading("Mapping từ yêu cầu sang thiết kế", level=2)
    add_table(
        doc,
        ["Lớp mô hình", "Vai trò", "Hiện thực trong VolunteerHub"],
        [
            ("Problem space", "Mô tả mục tiêu actor và use case không phụ thuộc công nghệ.",
             "Guest/Volunteer/Organizer/Sponsor/Admin, FR/UC, business rules."),
            ("Solution space", "Chuyển yêu cầu thành class, service, controller, DTO, UI và xử lý.",
             "EventService, RegistrationService, CertificateService, React pages, api.js facade."),
            ("Architectural space", "Quyết định cấu trúc triển khai, phân tầng, service và tích hợp.",
             "React/Vite, Ocelot Gateway, Auth/Event/Finance services, SQL Server shared database, SignalR."),
            ("Persistence space", "Thiết kế lưu trữ và mapping entity/table.",
             "MySqlDbContext, EF Core entities, indexes, foreign keys, migration, seed data."),
            ("Operational space", "Áp dụng NFR vào triển khai.",
             "JWT, IsActive middleware, rate limiting, audit log, export limit, responsive UI."),
        ],
        [2000, 3300, 4060],
    )
    doc.add_heading("UML artifact cần có cho đồ án", level=2)
    add_table(
        doc,
        ["Artifact", "Mục đích", "Nội dung áp dụng"],
        [
            ("Use case diagram", "Tổng quan actor và các nhóm mục tiêu.", "Nhóm public, volunteer, organizer, sponsor, admin."),
            ("Activity diagram", "Diễn tả business process hoặc luồng trong use case.", "Đăng ký event, check-in/out, campaign/donation, proposal."),
            ("Sequence diagram", "Thể hiện cộng tác runtime giữa UI, API, service, database, notification.", "Approve event, register, check-in, complete, donation confirm."),
            ("Class/domain model", "Mô tả entity và quan hệ nghiệp vụ bền vững.", "User, Event, Registration, WorkShift, Certificate, Campaign, Proposal."),
            ("State machine", "Mô tả vòng đời đối tượng có trạng thái.", "Event, Registration, Campaign, Donation, Proposal, Verification."),
            ("Component/deployment diagram", "Mô tả service, gateway, frontend, database và runtime port.", "BaseCore.WebClient, ApiGateway, AuthService, EventService, FinanceService."),
        ],
        [2200, 3100, 4060],
    )
    doc.add_heading("Nguyên tắc thiết kế áp dụng", level=2)
    add_bullets(
        doc,
        [
            "Tách controller khỏi business logic: controller nhận request, service chịu trách nhiệm nghiệp vụ.",
            "Tách persistence khỏi nghiệp vụ: entity/service không nên trộn lẫn mọi chi tiết lưu trữ nếu có repository/DbContext xử lý.",
            "Mỗi state transition quan trọng phải có guard condition, side effect và notification/audit rõ ràng.",
            "UI route guard chỉ hỗ trợ trải nghiệm; backend authorization mới là kiểm soát bắt buộc.",
            "Sequence diagram và API table phải khớp với code thật để tài liệu có thể dùng cho kiểm thử.",
            "NFR phải được gắn vào kiến trúc: bảo mật ở JWT/middleware, hiệu năng ở pagination/rate limit/export limit, maintainability ở service split.",
        ],
    )


SRS_USE_CASES = [
    {
        "id": "UC-01",
        "name": "Xem danh sách và chi tiết sự kiện công khai",
        "actor": "Guest",
        "goal": "Khách chưa đăng nhập có thể tìm hiểu các sự kiện đang mở hoặc đã hoàn thành để quyết định tham gia hoặc theo dõi tác động.",
        "pre": "Hệ thống có ít nhất một sự kiện Approved chưa kết thúc hoặc Completed.",
        "post": "Người dùng xem được thông tin sự kiện, organizer, thời gian, địa điểm, kỹ năng và tác động công khai nếu có.",
        "main": [
            "Guest truy cập landing page hoặc trang danh sách sự kiện.",
            "Hệ thống tải danh sách sự kiện public, mặc định chỉ hiển thị Approved chưa hết hạn và Completed.",
            "Guest lọc theo từ khóa, danh mục, kỹ năng hoặc vị trí nếu cần.",
            "Guest mở chi tiết sự kiện.",
            "Hệ thống hiển thị mô tả, thời gian, địa điểm, organizer, capacity, yêu cầu kỹ năng/KYC và thông tin tác động công khai.",
        ],
        "alt": [
            "Nếu không có sự kiện phù hợp, hệ thống hiển thị trạng thái rỗng và cho phép đổi bộ lọc.",
            "Nếu sự kiện đã Completed, hệ thống ưu tiên hiển thị impact thay vì nút đăng ký.",
        ],
        "rules": [
            "Guest không xem được Pending, Rejected hoặc Cancelled event.",
            "Approved event đã quá hạn không xuất hiện trong public listing mặc định.",
            "Thông tin tài trợ chỉ tính donation Confirmed và proposal Received/Reported.",
        ],
    },
    {
        "id": "UC-02",
        "name": "Xác thực chứng chỉ bằng mã verify",
        "actor": "Guest",
        "goal": "Bất kỳ người dùng nào cũng có thể kiểm chứng tính hợp lệ của chứng chỉ tình nguyện.",
        "pre": "Chứng chỉ đã được cấp sau khi event Completed.",
        "post": "Hệ thống trả thông tin chứng chỉ hoặc thông báo không tìm thấy.",
        "main": [
            "Guest nhập mã chứng chỉ hoặc truy cập link verify.",
            "Hệ thống tìm Certificate theo CertificateCode.",
            "Hệ thống hiển thị tên volunteer, event, ngày cấp và số giờ.",
            "Guest có thể tải PDF chứng chỉ nếu mã hợp lệ.",
        ],
        "alt": [
            "Nếu mã không tồn tại, hệ thống hiển thị thông báo chứng chỉ không hợp lệ.",
            "Nếu certificate đã bị thu hồi do admin uncomplete event, hệ thống không trả kết quả verify.",
        ],
        "rules": [
            "CertificateCode phải duy nhất.",
            "Verify certificate là public endpoint, không yêu cầu đăng nhập.",
            "PDF phải hiển thị đúng tiếng Việt.",
        ],
    },
    {
        "id": "UC-03",
        "name": "Đăng ký, đăng nhập và quản lý phiên",
        "actor": "Volunteer/Organizer/Sponsor/Admin",
        "goal": "Người dùng truy cập hệ thống theo đúng vai trò và duy trì phiên an toàn.",
        "pre": "Người dùng có thông tin đăng ký hợp lệ hoặc tài khoản đã được tạo.",
        "post": "Người dùng có JWT access token, refresh token và được điều hướng tới dashboard theo role.",
        "main": [
            "Người dùng đăng ký theo vai trò Volunteer, Organizer hoặc Sponsor.",
            "Hệ thống validate username/email/password và tạo tài khoản.",
            "Người dùng đăng nhập bằng định danh và mật khẩu.",
            "Backend xác thực, trả access token, refresh token và thông tin role.",
            "Frontend lưu phiên và gắn Authorization header khi gọi API.",
        ],
        "alt": [
            "Nếu access token hết hạn, interceptor gọi refresh token để lấy access token mới.",
            "Nếu refresh thất bại, frontend xóa phiên và chuyển về login.",
        ],
        "rules": [
            "User bị khóa IsActive=false bị chặn 401 trên service.",
            "Password và salt không được serialize về client.",
            "Route frontend và API backend đều phải kiểm role.",
        ],
    },
    {
        "id": "UC-04",
        "name": "Quản lý hồ sơ volunteer, KYC và kỹ năng",
        "actor": "Volunteer/Admin",
        "goal": "Volunteer khai báo năng lực và xác minh hồ sơ để tham gia các sự kiện yêu cầu điều kiện.",
        "pre": "Volunteer đã đăng nhập.",
        "post": "Hồ sơ, kỹ năng và trạng thái xác minh được cập nhật; admin có lịch sử xử lý.",
        "main": [
            "Volunteer cập nhật bio, avatar, nhóm máu, ngôn ngữ, sở thích và kỹ năng.",
            "Volunteer gửi minh chứng kỹ năng hoặc hồ sơ KYC.",
            "Admin xem danh sách hồ sơ chờ duyệt.",
            "Admin approve, reject hoặc request changes với lý do.",
            "Hệ thống gửi notification cho volunteer.",
        ],
        "alt": [
            "Nếu bị request changes, volunteer bổ sung minh chứng và gửi lại.",
            "Nếu bị reject, volunteer có thể gửi lại theo trạng thái cho phép.",
        ],
        "rules": [
            "KYC chỉ bắt buộc khi event yêu cầu.",
            "Reject/request changes phải có lý do rõ ràng tối thiểu 10 ký tự.",
            "VolunteerSkill unique theo UserId + SkillId.",
        ],
    },
    {
        "id": "UC-05",
        "name": "Organizer gửi xác minh tổ chức",
        "actor": "Organizer/Admin",
        "goal": "Đảm bảo organizer có hồ sơ tổ chức được kiểm duyệt trước khi event được phê duyệt.",
        "pre": "Organizer đã đăng nhập và tài khoản active.",
        "post": "OrganizerVerification ở trạng thái Pending, Verified, Rejected hoặc ChangesRequested.",
        "main": [
            "Organizer nhập tên tổ chức, đại diện, email, số điện thoại, địa chỉ, mô tả và giấy tờ.",
            "Hệ thống lưu hồ sơ ở trạng thái Pending.",
            "Admin xem hồ sơ trong màn verification.",
            "Admin approve nếu hợp lệ hoặc reject/request changes nếu chưa đạt.",
            "Hệ thống cập nhật trạng thái và thông báo organizer.",
        ],
        "alt": [
            "Nếu organizer sửa hồ sơ sau khi verified, hệ thống có thể yêu cầu duyệt lại theo policy.",
            "Admin ghi chú lý do để organizer biết cần bổ sung gì.",
        ],
        "rules": [
            "OrganizerId unique trong OrganizerVerification.",
            "Event chỉ được approve khi organizer đã Verified.",
            "Hồ sơ bị reject/request changes không làm mất lịch sử audit.",
        ],
    },
    {
        "id": "UC-06",
        "name": "Tạo, sửa và gửi duyệt sự kiện",
        "actor": "Organizer/Admin",
        "goal": "Organizer tạo sự kiện tình nguyện có đủ thông tin để admin kiểm duyệt và công khai.",
        "pre": "Organizer đã đăng nhập, active và có quyền tạo event.",
        "post": "Event được lưu ở Pending hoặc cập nhật trạng thái theo admin review.",
        "main": [
            "Organizer nhập title, mô tả, danh mục, kỹ năng, thời gian, địa điểm, tọa độ, max participants, ảnh và option KYC.",
            "Backend validate dữ liệu và tạo Event Pending.",
            "Admin duyệt event Pending.",
            "EventService kiểm organizer active/Verified, sinh QR và tạo channel.",
            "Event chuyển Approved và xuất hiện public.",
        ],
        "alt": [
            "Admin reject kèm lý do; organizer sửa và resubmit.",
            "Organizer sửa time/location của event Approved; hệ thống notify volunteer/sponsor liên quan.",
        ],
        "rules": [
            "Không approve event đã kết thúc.",
            "Reject reason tối thiểu 10 ký tự.",
            "Event có dữ liệu nghiệp vụ liên quan không xóa cứng; phải cancel.",
        ],
    },
    {
        "id": "UC-07",
        "name": "Đăng ký, rút đăng ký và xin hủy sau xác nhận",
        "actor": "Volunteer/Organizer",
        "goal": "Volunteer đăng ký tham gia đúng điều kiện và organizer kiểm soát danh sách tham gia.",
        "pre": "Event Approved, chưa bắt đầu, chưa hết chỗ.",
        "post": "Registration ở Pending, Confirmed hoặc Cancelled.",
        "main": [
            "Volunteer mở event Approved và gửi đăng ký.",
            "RegistrationService kiểm thời gian, capacity, KYC, shift và kỹ năng shift.",
            "Hệ thống tạo Registration Pending và thông báo organizer.",
            "Organizer xác nhận registration.",
            "Volunteer nhận thông báo đã được xác nhận.",
        ],
        "alt": [
            "Volunteer rút trực tiếp khi Pending.",
            "Volunteer đã Confirmed gửi cancel request; organizer duyệt hủy.",
            "Registration Cancelled có thể được đăng ký lại.",
        ],
        "rules": [
            "CurrentParticipants chỉ đếm Confirmed.",
            "Nếu event đã chia ca, registration bắt buộc có ShiftId.",
            "Volunteer không được đăng ký khi event đã bắt đầu hoặc đã kết thúc.",
        ],
    },
    {
        "id": "UC-08",
        "name": "Quản lý ca làm việc",
        "actor": "Organizer",
        "goal": "Chia nhỏ event thành các khung giờ/vị trí phù hợp khi sự kiện cần điều phối chi tiết.",
        "pre": "Event thuộc organizer và chưa có registration nếu bật/tạo ca mới.",
        "post": "WorkShift được tạo, cập nhật hoặc xóa theo ràng buộc.",
        "main": [
            "Organizer vào màn Manage Event, tab ca làm việc.",
            "Organizer tạo ca với tên, start/end, max volunteers và skill yêu cầu tùy chọn.",
            "Backend kiểm ca nằm trong khoảng thời gian event.",
            "Volunteer đăng ký event có ca phải chọn shift.",
        ],
        "alt": [
            "Organizer sửa ca khi chưa vi phạm dữ liệu đăng ký.",
            "Organizer xóa ca chỉ khi ca chưa có registration.",
        ],
        "rules": [
            "Không tạo ca sau khi event đã có registration.",
            "Shift phải thuộc event.",
            "Capacity shift tính cả Pending và Confirmed.",
        ],
    },
    {
        "id": "UC-09",
        "name": "Điểm danh QR/GPS, walk-in và check-out",
        "actor": "Volunteer/Organizer",
        "goal": "Ghi nhận sự có mặt và tính giờ tình nguyện thực tế.",
        "pre": "Registration Confirmed, event Approved và trong cửa sổ check-in.",
        "post": "Registration có AttendedAt, CheckedOutAt và VolunteerHours hợp lệ.",
        "main": [
            "Volunteer tự check-in hoặc organizer check-in cho volunteer.",
            "Hệ thống validate QR nếu event có QrCode; nếu không có QR thì kiểm GPS.",
            "Hệ thống ghi AttendedAt, đặt VolunteerHours = 0.",
            "Organizer check-out sau khi volunteer hoàn thành.",
            "Hệ thống tính giờ theo thời gian thực tế trong khung event/shift.",
        ],
        "alt": [
            "Organizer walk-in volunteer tại hiện trường.",
            "Organizer manual attend sau event trong cửa sổ 7 ngày.",
            "Organizer adjust hours trong giới hạn cho phép.",
        ],
        "rules": [
            "Shift check-in mở trước 15 phút và đóng sau 30 phút.",
            "Event check-in mở trước 30 phút và đóng sau 2 giờ.",
            "Check-in không cộng giờ ngay; check-out hoặc manual attend mới cập nhật giờ.",
        ],
    },
    {
        "id": "UC-10",
        "name": "Hoàn thành event, cấp chứng chỉ và huy hiệu",
        "actor": "Organizer/Admin/Volunteer",
        "goal": "Kết thúc event, ghi nhận đóng góp và phát hành chứng chỉ chính thống.",
        "pre": "Event ở trạng thái Approved.",
        "post": "Event Completed; certificate được cấp cho volunteer attended; badge được cập nhật nếu đủ điều kiện.",
        "main": [
            "Organizer hoặc Admin chọn hoàn thành event.",
            "Hệ thống đóng các registration Pending hoặc cancel request, không tính tham gia.",
            "CertificateService cấp chứng chỉ cho registration IsAttended=true.",
            "BadgeService kiểm điều kiện số sự kiện/số giờ.",
            "Volunteer xem certificate/badge trong Achievements hoặc Passport.",
        ],
        "alt": [
            "Admin uncomplete event, đưa event về Approved và thu hồi certificate.",
            "Admin auto-complete event quá hạn EndDate hơn 24 giờ.",
        ],
        "rules": [
            "Complete không phụ thuộc MinParticipants như điều kiện cứng.",
            "Certificate chỉ cấp cho attended registration.",
            "Rollback phải notify organizer.",
        ],
    },
    {
        "id": "UC-11",
        "name": "Trao đổi trong kênh sự kiện",
        "actor": "Volunteer/Organizer/Admin",
        "goal": "Tạo không gian trao đổi cho các thành viên liên quan đến event.",
        "pre": "Event đã Approved và channel đã được tạo.",
        "post": "Post/comment/like/poll được ghi nhận và cập nhật realtime khi có thể.",
        "main": [
            "Người dùng có quyền truy cập mở channel của event.",
            "Người dùng đăng post hoặc comment.",
            "ChannelService kiểm quyền và quan hệ cha-con.",
            "SignalRChannelRealtimeNotifier phát sự kiện realtime.",
        ],
        "alt": [
            "Organizer/Admin pin hoặc xóa nội dung không phù hợp.",
            "Người dùng tạo poll để lấy ý kiến thành viên.",
        ],
        "rules": [
            "Chỉ volunteer confirmed, organizer sở hữu hoặc admin được truy cập channel private.",
            "Comment phải thuộc post trong cùng channel.",
            "Like unique theo PostId + UserId.",
        ],
    },
    {
        "id": "UC-12",
        "name": "Campaign ủng hộ cá nhân và donation",
        "actor": "Organizer/Volunteer/Admin",
        "goal": "Ghi nhận minh bạch các khoản ủng hộ cá nhân cho event.",
        "pre": "Event Approved và organizer tạo campaign hợp lệ.",
        "post": "Donation được xác nhận/từ chối; campaign có thể đóng và báo cáo sử dụng.",
        "main": [
            "Organizer tạo campaign Draft, mở Open khi sẵn sàng.",
            "Volunteer gửi donation với số tiền, display name, proof image và tùy chọn ẩn danh.",
            "Organizer đối soát ngoài hệ thống và confirm/reject.",
            "Hệ thống chỉ cộng Confirmed donation vào tổng public.",
            "Organizer đóng campaign và báo cáo sử dụng tiền.",
        ],
        "alt": [
            "User hủy donation khi còn PendingConfirmation.",
            "Event hủy làm Pending donation tự Cancelled.",
            "Admin finance watch phát hiện donation pending quá hạn.",
        ],
        "rules": [
            "Số tiền donation phải lớn hơn 0.",
            "UsedAmount không vượt ConfirmedAmount nếu không có giải trình.",
            "Anonymous donation không hiển thị thông tin liên hệ public.",
        ],
    },
    {
        "id": "UC-13",
        "name": "Tài trợ doanh nghiệp bằng sponsorship proposal",
        "actor": "Organizer/Sponsor/Admin",
        "goal": "Quản lý tài trợ doanh nghiệp hai chiều và minh bạch trạng thái nhận tiền.",
        "pre": "Sponsor active, event Approved và chưa có proposal active trùng cho sponsor/event.",
        "post": "Proposal chuyển trạng thái theo state machine và được tính impact khi Received/Reported.",
        "main": [
            "Organizer gửi lời mời tài trợ hoặc Sponsor gửi đề nghị tài trợ.",
            "Bên nhận accept/reject proposal.",
            "Organizer mark received khi đã nhận tiền/thỏa thuận hiện vật.",
            "Organizer report sử dụng tài trợ sau event.",
            "Sponsor theo dõi proposal trong My Sponsorships.",
        ],
        "alt": [
            "Bên tạo cancel khi proposal còn Pending.",
            "Admin revert Accepted về Pending khi cần xử lý sai sót.",
            "Event hủy tự cancel proposal Pending/Accepted.",
        ],
        "rules": [
            "Mỗi sponsor chỉ có một proposal active trên một event.",
            "Sponsor không hủy sau Accepted.",
            "Chỉ Received/Reported tính vào public impact.",
        ],
    },
    {
        "id": "UC-14",
        "name": "Đánh giá hai chiều và moderation",
        "actor": "Volunteer/Organizer/Admin",
        "goal": "Ghi nhận uy tín sau sự kiện và cho phép admin xử lý nội dung không phù hợp.",
        "pre": "Event Completed và actor có quan hệ tham gia/sở hữu phù hợp.",
        "post": "Rating được tạo, ẩn/hiện/xóa theo quyền admin nếu cần.",
        "main": [
            "Volunteer đã attended đánh giá organizer.",
            "Organizer đánh giá volunteer đã attended.",
            "Hệ thống validate mỗi cặp chỉ đánh giá một lần trong một event.",
            "Rating public hiển thị trên profile khi không bị ẩn.",
        ],
        "alt": [
            "Admin hide rating kèm lý do.",
            "Admin unhide hoặc delete rating nếu vi phạm.",
        ],
        "rules": [
            "Score nằm trong khoảng 1-5.",
            "Unique theo EventId + RaterId + RateeId.",
            "Delete rating chỉ Admin được thực hiện.",
        ],
    },
    {
        "id": "UC-15",
        "name": "Quản trị hệ thống, giám sát và export",
        "actor": "Admin",
        "goal": "Admin vận hành hệ thống an toàn, truy vết được và không phá vỡ dữ liệu lịch sử.",
        "pre": "Admin đã đăng nhập, tài khoản active.",
        "post": "Dữ liệu được duyệt/cập nhật/xuất báo cáo theo đúng quyền và audit.",
        "main": [
            "Admin xem dashboard, monitoring và audit logs.",
            "Admin quản lý user, khóa/mở tài khoản và xem impact summary.",
            "Admin duyệt event/KYC/skill/organizer verification.",
            "Admin quản lý category, skill, badge và rating.",
            "Admin export events/users/finance dạng JSON hoặc CSV.",
        ],
        "alt": [
            "Khóa organizer tự cascade cancel event active theo quy tắc.",
            "Xóa cứng bị chặn nếu có dữ liệu nghiệp vụ liên quan.",
            "Finance watch chỉ đọc để phát hiện donation/proposal/campaign cần đối soát.",
        ],
        "rules": [
            "Export có maxRows limit và chống CSV injection.",
            "Audit log/monitoring/export không sửa/xóa qua UI demo.",
            "Các thao tác admin nhạy cảm cần ghi audit.",
        ],
    },
]


SRS_FR_DETAILS = [
    ("FR-01", "Xác thực, phân quyền và quản lý phiên", "Tất cả user",
     "Hệ thống cung cấp đăng ký, đăng nhập, JWT access token, refresh token và kiểm soát quyền truy cập theo vai trò.",
     ["Thông tin đăng ký/đăng nhập", "Role người dùng", "Refresh token khi làm mới phiên"],
     ["Access token", "Refresh token", "Thông tin user gồm id, name, email, role"],
     ["Đăng nhập sai trả lỗi rõ ràng.", "User inactive bị chặn 401.", "Frontend tự refresh access token khi còn refresh token hợp lệ."]),
    ("FR-02", "Trang công khai và khám phá sự kiện", "Guest, Volunteer, Organizer, Sponsor, Admin",
     "Người dùng xem landing page, danh sách sự kiện công khai, bản đồ và chi tiết sự kiện.",
     ["Keyword", "CategoryId", "SkillId", "Location", "Page/PageSize"],
     ["Danh sách event", "Thông tin event detail", "Impact public nếu có"],
     ["Không lộ Pending/Rejected/Cancelled trên public.", "Filter hoạt động độc lập và kết hợp.", "Trang chi tiết hiển thị rõ trạng thái đăng ký."]),
    ("FR-03", "Hồ sơ tình nguyện viên và kỹ năng", "Volunteer, Admin",
     "Volunteer quản lý thông tin cá nhân, kỹ năng, minh chứng và trạng thái xác minh kỹ năng.",
     ["Bio/avatar", "Danh sách kỹ năng", "Evidence URL", "Level"],
     ["VolunteerProfile", "VolunteerSkill", "VerificationStatus"],
     ["Không cho trùng skill trên cùng user.", "Admin approve/reject/request changes được từng kỹ năng.", "Profile hiển thị được passport summary."]),
    ("FR-04", "Volunteer Passport", "Volunteer",
     "Hệ thống tổng hợp lịch sử tham gia, giờ tình nguyện, chứng chỉ và huy hiệu của volunteer.",
     ["UserId từ token", "Registration history", "Certificate", "UserBadge"],
     ["Danh sách hoạt động", "Tổng giờ", "Chứng chỉ", "Huy hiệu"],
     ["Chỉ registration attended mới tính giờ.", "Giờ thay đổi sau adjust/check-out phải cập nhật passport.", "Certificate verify code có thể mở từ passport."]),
    ("FR-05", "KYC tình nguyện viên", "Volunteer, Admin",
     "Volunteer gửi ảnh giấy tờ và admin xử lý để cho phép tham gia event yêu cầu KYC.",
     ["IdentityFrontImageUrl", "IdentityBackImageUrl", "PortraitImageUrl", "AdminNote"],
     ["KycStatus", "Notification", "Audit trail"],
     ["Event RequiresKyc chặn volunteer chưa Verified.", "Reject/request changes có lý do tối thiểu 10 ký tự.", "Volunteer gửi lại được khi Rejected/ChangesRequested."]),
    ("FR-06", "Xác minh tổ chức", "Organizer, Admin",
     "Organizer gửi hồ sơ pháp lý/tổ chức; admin duyệt trước khi sự kiện được approve.",
     ["OrganizationName", "RepresentativeName", "ContactEmail", "Phone", "Address", "DocumentUrl"],
     ["OrganizerVerification.Status", "VerifiedAt", "VerifiedBy", "Notification"],
     ["Organizer chưa Verified không được admin approve event.", "Mỗi organizer có một hồ sơ verification.", "Admin có thể request changes thay vì reject ngay."]),
    ("FR-07", "Quản lý vòng đời sự kiện", "Organizer, Admin",
     "Organizer tạo/sửa event; admin duyệt/từ chối/hủy/chuyển quyền/hoàn thành hoặc mở lại event.",
     ["EventCreateDto", "EventUpdateDto", "Reject/Cancel reason", "Transfer organizer id"],
     ["Event status", "QrCode", "Channel", "Notifications"],
     ["Approve sinh QR và channel.", "Cancel cascade campaign/proposal/donation pending.", "Không xóa cứng event có dữ liệu nghiệp vụ."]),
    ("FR-08", "Đăng ký tham gia sự kiện", "Volunteer",
     "Volunteer đăng ký event Approved với kiểm tra capacity, KYC, shift và kỹ năng.",
     ["EventId", "ShiftId optional", "Note"],
     ["Registration Pending", "Notification cho organizer"],
     ["Chặn đăng ký khi event đã bắt đầu/kết thúc.", "Event có ca thì ShiftId bắt buộc.", "Registration active không được trùng."]),
    ("FR-09", "Quản lý đăng ký bởi organizer", "Organizer, Admin",
     "Organizer xác nhận/hủy/chuyển ca/xử lý yêu cầu hủy và xem danh sách đăng ký.",
     ["RegistrationId", "Action", "Cancel reason", "NewShiftId"],
     ["Registration status", "CurrentParticipants", "Notification"],
     ["Confirm chỉ từ Pending.", "Cancel Confirmed phải giảm participant count.", "Change shift không cho sau khi attended."]),
    ("FR-10", "Quản lý ca làm việc", "Organizer",
     "Organizer tạo/sửa/xóa ca làm việc khi sự kiện cần chia ca.",
     ["Name", "StartTime", "EndTime", "MaxVolunteers", "RequiredSkillId"],
     ["WorkShift", "Shift channel nếu có", "Registration shift binding"],
     ["Shift nằm trong thời gian event.", "Không tạo ca sau khi có registration.", "Xóa ca bị chặn nếu đã có registration."]),
    ("FR-11", "Điểm danh, walk-in, manual attend và check-out", "Volunteer, Organizer",
     "Hệ thống ghi nhận tham gia bằng QR/GPS và tính giờ khi check-out hoặc manual attend.",
     ["QrCode", "Latitude/Longitude", "RegistrationId", "Hours override"],
     ["AttendedAt", "CheckedOutAt", "VolunteerHours", "Audit metadata"],
     ["Check-in đúng cửa sổ thời gian.", "QR sai hoặc GPS ngoài bán kính bị chặn.", "Check-out tính giờ không vượt khung event/shift."]),
    ("FR-12", "Hoàn thành sự kiện và impact", "Organizer, Admin, Guest",
     "Event Completed tạo báo cáo tác động công khai và đóng các registration chưa xử lý.",
     ["EventId", "ManualCompletionAttendance optional"],
     ["Event Completed", "Impact metrics", "Notifications"],
     ["Pending registration không tính tham gia.", "Impact chỉ tính dữ liệu hợp lệ.", "Admin có overdue preview/auto-complete."]),
    ("FR-13", "Chứng chỉ điện tử và huy hiệu", "Volunteer, Guest, Admin",
     "Cấp chứng chỉ PDF có mã verify và trao huy hiệu tự động theo thành tích.",
     ["Registration attended", "VolunteerHours", "Badge condition"],
     ["Certificate PDF", "CertificateCode", "UserBadge"],
     ["Certificate unique theo user/event.", "Guest verify được mã chứng chỉ.", "Uncomplete thu hồi certificate."]),
    ("FR-14", "Đánh giá hai chiều", "Volunteer, Organizer, Admin",
     "Sau event Completed, các bên đánh giá lẫn nhau và admin moderation nội dung.",
     ["EventId", "RateeId", "Score", "Comment"],
     ["Rating", "Public rating list", "Hidden state"],
     ["Mỗi cặp chỉ rating một lần/event.", "Score 1-5.", "Admin hide/unhide/delete rating."]),
    ("FR-15", "Channel trao đổi sự kiện", "Volunteer, Organizer, Admin",
     "Kênh trao đổi được tạo khi event Approved, hỗ trợ post/comment/like/poll và realtime.",
     ["ChannelId", "Post content", "Comment", "Poll options", "Attachment"],
     ["Post", "Comment", "Like", "Poll vote", "Realtime event"],
     ["Kiểm quyền truy cập channel.", "Quan hệ channel-post-comment phải đúng.", "SignalR cập nhật realtime khi khả dụng."]),
    ("FR-16", "Campaign ủng hộ cá nhân", "Organizer, Volunteer, Admin",
     "Organizer tạo đợt kêu gọi, volunteer gửi donation, organizer xác nhận và báo cáo sử dụng.",
     ["Campaign data", "Donation amount", "Proof image", "Report summary"],
     ["Campaign status", "Donation status", "Confirmed amount", "Report"],
     ["Donation chỉ tính khi Confirmed.", "Campaign theo Draft/Open/Closed/Reported.", "Admin finance watch phát hiện quá hạn."]),
    ("FR-17", "Tài trợ doanh nghiệp", "Organizer, Sponsor, Admin",
     "Quản lý proposal hai chiều giữa organizer và sponsor, ghi nhận nhận tiền và báo cáo.",
     ["OrganizerRequest", "SponsorOffer", "ResponseMessage", "ActualReceivedAmount", "Report"],
     ["Proposal status", "Sponsor tracking", "Impact sponsorship amount"],
     ["Chống duplicate active proposal.", "Sponsor không hủy sau Accepted.", "Received mới tính vào impact."]),
    ("FR-18", "Dashboard và thông báo", "Tất cả user đã đăng nhập",
     "Dashboard tùy role và notification điều hướng người dùng tới việc cần xử lý.",
     ["User role", "Notification type", "RelatedId"],
     ["Dashboard metrics", "Notification list", "Read state"],
     ["Dashboard phản ánh role.", "Notification phân trang và mark read.", "Link notification dẫn đúng màn hình."]),
    ("FR-19", "Admin vận hành hệ thống", "Admin",
     "Admin quản trị user, verification, danh mục, huy hiệu, rating, monitoring, finance watch và export.",
     ["Filter/search", "Review action", "Export format", "Toggle status"],
     ["Danh sách quản trị", "Audit log", "CSV/JSON export", "Impact summary"],
     ["Export giới hạn maxRows.", "Không sửa/xóa finance trực tiếp qua watch.", "Khóa organizer cascade đúng policy."]),
    ("FR-20", "Upload file và ảnh", "Authenticated user",
     "Người dùng upload ảnh/file phục vụ event, KYC, skill evidence và donation proof.",
     ["Multipart file", "File type", "Context sử dụng"],
     ["File URL", "Upload response"],
     ["Chỉ trả URL hợp lệ.", "Frontend dùng URL lưu vào DTO tương ứng.", "Upload lỗi phải có thông báo rõ."]),
    ("FR-21", "Gợi ý và knowledge graph", "Volunteer, Organizer, Admin",
     "Hệ thống gợi ý event/volunteer/sponsor theo kỹ năng và quan hệ trong graph khi khả dụng.",
     ["UserId", "EventId", "Limit"],
     ["Recommended events", "Recommended volunteers/sponsors", "Degraded flag"],
     ["Có fallback khi graph không khả dụng.", "Không gợi ý event không public.", "Kết quả giới hạn theo limit."]),
]


def add_srs_functional_detail(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("ĐẶC TẢ YÊU CẦU CHỨC NĂNG CHI TIẾT", level=1)
    doc.add_paragraph(
        "Phần này diễn giải từng yêu cầu chức năng ở mức đủ để nhóm phát triển, kiểm thử và người đọc đồ án hiểu rõ phạm vi, "
        "dữ liệu vào/ra, ràng buộc và tiêu chí nghiệm thu."
    )
    for code, name, actors, desc, inputs, outputs, acceptance in SRS_FR_DETAILS:
        add_page_section(doc, f"{code} - {name}", level=2)
        add_table(
            doc,
            ["Mục", "Nội dung"],
            [
                ("Actor liên quan", actors),
                ("Mô tả", desc),
                ("Ưu tiên", "Cao - thuộc phạm vi chính của đồ án VolunteerHub."),
                ("Nguồn đối chiếu", "docs/2-yeu-cau-chuc-nang.md, source controller/service/frontend tương ứng."),
            ],
            [2100, 7260],
        )
        doc.add_heading("Dữ liệu vào", level=3)
        add_bullets(doc, inputs)
        doc.add_heading("Dữ liệu ra", level=3)
        add_bullets(doc, outputs)
        doc.add_heading("Tiêu chí nghiệm thu", level=3)
        add_bullets(doc, acceptance)
        doc.add_heading("Ghi chú kiểm thử", level=3)
        add_bullets(
            doc,
            [
                "Kiểm thử trường hợp thành công và tối thiểu một trường hợp lỗi nghiệp vụ.",
                "Kiểm thử phân quyền theo role và user inactive nếu yêu cầu có API bảo vệ.",
                "Kiểm thử dữ liệu hiển thị ở frontend sau khi backend thay đổi trạng thái.",
            ],
        )


def add_detailed_srs_appendix(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("ĐẶC TẢ USE CASE CHI TIẾT", level=1)
    doc.add_paragraph(
        "Phần này mở rộng đặc tả use case theo mẫu CNPM: mỗi use case có mục tiêu, actor, tiền điều kiện, "
        "hậu điều kiện, luồng chính, luồng thay thế và quy tắc nghiệp vụ liên quan."
    )
    for uc in SRS_USE_CASES:
        add_page_section(doc, f"{uc['id']} - {uc['name']}", level=2)
        add_table(
            doc,
            ["Mục", "Nội dung"],
            [
                ("Actor chính", uc["actor"]),
                ("Mục tiêu", uc["goal"]),
                ("Tiền điều kiện", uc["pre"]),
                ("Hậu điều kiện", uc["post"]),
            ],
            [2100, 7260],
        )
        doc.add_heading("Luồng chính", level=3)
        add_numbered(doc, uc["main"])
        doc.add_heading("Luồng thay thế/ngoại lệ", level=3)
        add_bullets(doc, uc["alt"])
        doc.add_heading("Quy tắc nghiệp vụ", level=3)
        add_bullets(doc, uc["rules"])
        doc.add_heading("Dữ liệu/API liên quan", level=3)
        related = {
            "Guest": "EventsController, CertificatesController, bảng Events/Certificates.",
            "Volunteer": "ProfileController, RegistrationsController, SupportCampaignController, bảng VolunteerProfiles/Registrations/IndividualDonations.",
            "Organizer": "EventsController, WorkShiftsController, RegistrationsController, SupportCampaignController, bảng Events/WorkShifts/Registrations/SupportCampaigns.",
            "Organizer/Admin": "EventsController, EventService, OrganizerVerificationController, bảng Events/OrganizerVerifications/Channels.",
            "Organizer/Volunteer": "RegistrationsController, RegistrationService, bảng Registrations/WorkShifts/Notifications.",
            "Volunteer/Organizer": "RegistrationsController, RegistrationService, bảng Registrations/Events/WorkShifts.",
            "Volunteer/Organizer/Admin": "ChannelsController hoặc RatingsController, bảng Channels/Ratings.",
            "Organizer/Volunteer/Admin": "SupportCampaignController, AdminFinanceController, bảng SupportCampaigns/IndividualDonations.",
            "Organizer/Sponsor/Admin": "SponsorshipProposalController, SponsorProfileController, bảng SponsorshipProposals/SponsorProfiles.",
            "Admin": "AdminController, MonitoringController, AdminFinanceController, bảng Users/AuditLogs/Notifications.",
        }
        doc.add_paragraph(related.get(uc["actor"], "Các controller/service tương ứng với module nghiệp vụ."))

    doc.add_page_break()
    doc.add_heading("MA TRẬN TRUY VẾT YÊU CẦU", level=1)
    trace_rows = [
        ("FR-01", "UC-03", "AuthController, UserController", "Login/Register/Refresh, ProtectedRoute"),
        ("FR-02", "UC-01", "EventsController", "LandingPage, EventList, EventDetail"),
        ("FR-03/04/05", "UC-04", "ProfileController, AdminController", "Profile, AdminVerifications"),
        ("FR-06", "UC-05", "OrganizerVerificationController", "OrganizerVerification, AdminVerifications"),
        ("FR-07", "UC-06", "EventsController, EventService", "EventForm, AdminEvents, MyEvents"),
        ("FR-08/09", "UC-07", "RegistrationsController, RegistrationService", "EventDetail, ManageEvent/RegistrationsTab"),
        ("FR-10", "UC-08", "WorkShiftsController", "ManageEvent/ShiftsTab"),
        ("FR-11", "UC-09", "RegistrationsController, RegistrationService", "CheckInTab, VolunteerCheckInModal"),
        ("FR-12/13", "UC-10", "EventService, CertificateService, BadgeService", "Achievements, VerifyCertificate"),
        ("FR-15", "UC-11", "ChannelsController, ChannelService, ChannelHub", "Channel"),
        ("FR-16", "UC-12", "SupportCampaignController", "CampaignsTab, Activity donations"),
        ("FR-17", "UC-13", "SponsorshipProposalController", "CorporateTab, MySponsorships"),
        ("FR-14", "UC-14", "RatingsController", "Activity, AdminRatings"),
        ("FR-18/19", "UC-15", "AdminController, MonitoringController", "Admin pages, Dashboard, Notifications"),
    ]
    add_table(doc, ["Yêu cầu", "Use case", "Backend", "Frontend"], trace_rows, [1200, 1200, 3300, 3660])

    doc.add_heading("YÊU CẦU VALIDATION CHI TIẾT", level=1)
    validation_rows = [
        ("Đăng ký tài khoản", "Username/email/password hợp lệ; role thuộc tập cho phép.", "400 nếu trùng hoặc thiếu dữ liệu."),
        ("Tạo event", "Title, thời gian, địa điểm, max participants hợp lệ; StartDate < EndDate.", "400 nếu thiếu hoặc sai logic."),
        ("Approve event", "Event Pending, chưa kết thúc, organizer active và Verified.", "Không approve event quá hạn hoặc organizer chưa xác minh."),
        ("Register event", "Event Approved, chưa bắt đầu/kết thúc, chưa full, đủ KYC/shift/skill.", "Không tạo duplicate registration active."),
        ("Create shift", "Shift thuộc event, nằm trong StartDate/EndDate, chưa có registration.", "Chặn tạo/sửa/xóa gây sai lịch sử."),
        ("Check-in", "Registration Confirmed, QR đúng hoặc GPS trong bán kính, đúng cửa sổ thời gian.", "Không check-in hai lần."),
        ("Check-out", "Đã check-in, chưa check-out, thời điểm check-out >= AttendedAt.", "VolunteerHours tính trong khung event/shift."),
        ("Donation", "Amount > 0, campaign Open, display name nếu không anonymous.", "Chỉ Confirmed tính public."),
        ("Proposal", "Không có proposal active trùng sponsor/event.", "Sponsor không cancel sau Accepted."),
        ("Admin delete", "Chỉ xóa cứng khi chưa có dữ liệu nghiệp vụ liên quan.", "Nếu có lịch sử, dùng cancel/hide/status."),
    ]
    add_table(doc, ["Nghiệp vụ", "Điều kiện hợp lệ", "Xử lý lỗi"], validation_rows, [2200, 4300, 2860])

    doc.add_heading("KỊCH BẢN NGHIỆP VỤ NGHIỆM THU", level=1)
    scenarios = [
        ("SC-01", "Organizer verified tạo event không chia ca", "Event Pending -> Approved; Volunteer đăng ký không cần ShiftId."),
        ("SC-02", "Organizer tạo event có chia ca", "Volunteer bắt buộc chọn shift; shift capacity được kiểm tra."),
        ("SC-03", "Volunteer chưa KYC đăng ký event yêu cầu KYC", "Hệ thống chặn và yêu cầu hoàn tất KYC."),
        ("SC-04", "Volunteer đăng ký, organizer xác nhận", "Registration Pending -> Confirmed, CurrentParticipants tăng."),
        ("SC-05", "Volunteer xin hủy sau Confirmed", "CancelRequested=true; organizer xử lý sang Cancelled."),
        ("SC-06", "Check-in QR đúng", "IsAttended=true, AttendedAt có giá trị, VolunteerHours vẫn là 0."),
        ("SC-07", "Check-out sau check-in", "CheckedOutAt có giá trị, VolunteerHours tính theo thời gian thực tế."),
        ("SC-08", "Manual attend sau event", "Cho phép trong 7 ngày sau EndDate, cập nhật giờ hợp lệ."),
        ("SC-09", "Complete event còn Pending registration", "Pending bị Cancelled và không cấp chứng chỉ."),
        ("SC-10", "Admin uncomplete event", "Event về Approved, certificate bị thu hồi."),
        ("SC-11", "Donation pending quá hạn", "Admin finance watch phát hiện để organizer đối soát."),
        ("SC-12", "Sponsor offer và organizer nhận tiền", "Proposal Pending -> Accepted -> Received, impact cập nhật."),
        ("SC-13", "Event bị hủy", "Campaign Open/Draft đóng, pending donation/proposal active bị cancel."),
        ("SC-14", "Rating sau event", "Chỉ người có quan hệ tham gia/sở hữu được rating."),
        ("SC-15", "Khóa organizer", "Event active của organizer bị cascade cancel theo policy."),
    ]
    add_table(doc, ["Mã", "Kịch bản", "Kết quả mong đợi"], scenarios, [900, 3600, 4860])


DETAIL_MODULES = [
    ("Identity/Auth", "Đăng ký, đăng nhập, refresh token, user active middleware, role mapping.", [
        "AuthController nhận request đăng ký/đăng nhập.",
        "TokenHelper sinh JWT chứa user id và role.",
        "AuthRefreshToken lưu hash refresh token để quản lý phiên.",
        "Frontend api.js tự gắn bearer token và refresh khi 401.",
    ]),
    ("Profile/KYC/Skill", "Quản lý hồ sơ volunteer, kỹ năng và các trạng thái xác minh.", [
        "ProfileController đọc/cập nhật VolunteerProfile.",
        "VolunteerSkill lưu level, evidenceUrl và verificationStatus.",
        "AdminController xử lý approve/reject/request changes.",
        "NotificationService gửi kết quả xét duyệt.",
    ]),
    ("Organizer Verification", "Xác minh tổ chức trước khi sự kiện được duyệt.", [
        "OrganizerVerificationController lưu hồ sơ theo OrganizerId.",
        "Admin review cập nhật Status, AdminNote, VerifiedAt, VerifiedBy.",
        "EventService.ApproveAsync kiểm organizer Verified trước khi approve.",
    ]),
    ("Event Lifecycle", "Vòng đời event và các cascade liên quan.", [
        "CreateAsync đặt Status=Pending.",
        "ApproveAsync sinh QrCode và tạo Channel.",
        "CancelAsync đóng campaign, hủy pending donation/proposal active và notify.",
        "CompleteAsync cấp certificate và xử lý registration chưa hoàn tất.",
    ]),
    ("Registration/Attendance", "Đăng ký, xác nhận, điểm danh và tính giờ.", [
        "RegisterAsync dùng transaction Serializable và UPDLOCK để tránh vượt capacity.",
        "ConfirmAsync kiểm capacity và interview gate nếu event yêu cầu.",
        "ValidateCheckInWindow kiểm cửa sổ event/shift.",
        "CalculateActualVolunteerHours tính giờ từ AttendedAt đến CheckedOutAt trong khung hợp lệ.",
    ]),
    ("Finance", "Campaign, donation và sponsorship proposal.", [
        "SupportCampaignController quản lý state Draft/Open/Closed/Reported.",
        "IndividualDonation chỉ cộng public khi Confirmed.",
        "SponsorshipProposalController chống duplicate active proposal.",
        "AdminFinanceController chỉ đọc dữ liệu đối soát, không bypass workflow.",
    ]),
    ("Engagement", "Channel, notification, rating, certificate, badge.", [
        "ChannelService kiểm quyền truy cập channel theo event/registration/admin.",
        "ChannelHub phát realtime cho post/comment.",
        "RatingController validate quan hệ rater/ratee/event.",
        "BadgeService kiểm điều kiện sau khi volunteer hours thay đổi.",
    ]),
    ("Admin Operations", "Quản trị người dùng, danh mục, giám sát và export.", [
        "AdminController gom user, verification, export và impact summary.",
        "MonitoringController trả health, summary và audit logs.",
        "Delete operations bị chặn nếu có dữ liệu nghiệp vụ liên quan.",
        "Export có giới hạn maxRows và chống CSV injection.",
    ]),
]


def add_detail_deep_content(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("THIẾT KẾ MODULE CHI TIẾT", level=1)
    for name, desc, steps in DETAIL_MODULES:
        add_page_section(doc, name, level=2)
        doc.add_paragraph(desc)
        doc.add_heading("Thành phần xử lý", level=3)
        add_numbered(doc, steps)
        doc.add_heading("Input/Output chính", level=3)
        rows = [
            ("Input", "Request DTO từ frontend hoặc route parameter; JWT claims để xác định user/role."),
            ("Processing", "Controller validate sơ bộ, service kiểm nghiệp vụ, DbContext cập nhật entity, notification/audit nếu cần."),
            ("Output", "JSON response theo trạng thái thành công/lỗi; frontend hiển thị toast, modal hoặc chuyển route."),
            ("Error", "Exception nghiệp vụ được chuyển thành response lỗi, không làm thay đổi dữ liệu nếu transaction rollback."),
        ]
        add_table(doc, ["Loại", "Mô tả"], rows, [1800, 7560])
        doc.add_heading("Rủi ro và kiểm soát", level=3)
        add_bullets(
            doc,
            [
                "Luôn kiểm quyền ở backend, không tin vào route guard frontend.",
                "Không xóa dữ liệu đã có lịch sử; dùng trạng thái Cancelled/Hidden/Inactive.",
                "Các thao tác thay đổi trạng thái cần notification để người dùng biết kết quả.",
                "Các danh sách lớn cần pagination/filter để tránh tải quá nhiều dữ liệu.",
            ],
        )

    doc.add_page_break()
    doc.add_heading("ĐẶC TẢ BẢNG DỮ LIỆU CHI TIẾT", level=1)
    entities = [
        ("Users", [("Id", "int PK"), ("UserName", "string unique"), ("Password/Salt", "ẩn khỏi JSON"), ("UserType", "0 Volunteer, 1 Organizer, 2 Sponsor, 3 Admin"), ("IsActive", "chặn user bị khóa")]),
        ("Events", [("Id", "int PK"), ("OrganizerId", "FK Users"), ("CategoryId", "FK EventCategories"), ("Status", "Pending/Approved/Completed/Rejected/Cancelled"), ("QrCode", "GUID-based khi approve"), ("RequiredSkillIds", "JSON array")]),
        ("Registrations", [("EventId/UserId", "unique"), ("ShiftId", "nullable FK WorkShifts"), ("Status", "Pending/Confirmed/Cancelled"), ("AttendedAt/CheckedOutAt", "mốc check-in/out"), ("VolunteerHours", "decimal pro-rate")]),
        ("WorkShifts", [("EventId", "FK Events"), ("StartTime/EndTime", "nằm trong event"), ("MaxVolunteers", "capacity ca"), ("RequiredSkillId", "nullable FK Skills")]),
        ("VolunteerProfiles", [("UserId", "unique"), ("KycStatus", "Unverified/Pending/Verified/Rejected/ChangesRequested"), ("Identity images", "minh chứng KYC"), ("TotalVolunteerHours", "tổng hợp từ registrations")]),
        ("VolunteerSkills", [("UserId/SkillId", "unique"), ("Level", "mức kỹ năng"), ("VerificationStatus", "SelfDeclared/Pending/Verified/Rejected/ChangesRequested"), ("EvidenceUrl", "minh chứng")]),
        ("OrganizerVerifications", [("OrganizerId", "unique"), ("Status", "Pending/Verified/Rejected/ChangesRequested"), ("DocumentUrl", "giấy tờ"), ("VerifiedBy", "admin reviewer")]),
        ("SupportCampaigns", [("EventId", "FK Events"), ("Status", "Draft/Open/Closed/Cancelled/Reported"), ("TargetAmount", "mục tiêu"), ("UsedAmount/ReportSummary", "báo cáo sử dụng")]),
        ("IndividualDonations", [("CampaignId", "FK SupportCampaigns"), ("UserId", "donor"), ("Amount", "> 0"), ("Status", "PendingConfirmation/Confirmed/Rejected/Cancelled"), ("ProofImageUrl", "minh chứng")]),
        ("SponsorshipProposals", [("EventId/SponsorId/OrganizerId", "quan hệ tài trợ"), ("Type", "OrganizerRequest/SponsorOffer"), ("Status", "Pending/Accepted/Received/Reported/Rejected/Cancelled"), ("ActualReceivedAmount", "tiền đã nhận")]),
        ("Certificates", [("CertificateCode", "unique"), ("UserId/EventId", "unique"), ("VolunteerHours", "giờ trên chứng chỉ"), ("PdfUrl", "đường dẫn nếu lưu file")]),
        ("Channels/Posts/Comments/Likes", [("Channel.EventId", "FK Events"), ("Post.AuthorId", "FK Users"), ("Comment.PostId", "FK Posts"), ("Like.PostId/UserId", "unique")]),
        ("Ratings", [("EventId", "FK Events"), ("RaterId/RateeId", "FK Users"), ("Score", "1-5"), ("IsHidden", "moderation"), ("EventId/RaterId/RateeId", "unique")]),
        ("AuditLogs/Notifications", [("AuditLogs.Metadata", "JSON metadata"), ("Notifications.UserId", "người nhận"), ("Type/RelatedId", "điều hướng frontend"), ("IsRead", "trạng thái đọc")]),
    ]
    for entity, fields in entities:
        add_page_section(doc, entity, level=2)
        add_table(doc, ["Trường", "Thiết kế/Ghi chú"], fields, [2600, 6760])
        doc.add_heading("Quan hệ và lưu ý", level=3)
        add_bullets(
            doc,
            [
                "Entity được cấu hình trong MySqlDbContext bằng Fluent API để xác định khóa, index và quan hệ.",
                "Các quan hệ nghiệp vụ quan trọng được validate thêm ở service/controller, không chỉ dựa vào database.",
                "Các trường trạng thái dùng string để dễ đọc trong đồ án và tương thích frontend.",
            ],
        )

    doc.add_page_break()
    doc.add_heading("ĐẶC TẢ API THEO NHÓM", level=1)
    api_groups = [
        ("Auth/Profile", [
            ("POST /api/auth/register", "Guest", "Tạo tài khoản."),
            ("POST /api/auth/login", "Guest", "Đăng nhập."),
            ("POST /api/auth/refresh", "Authenticated", "Làm mới access token."),
            ("GET/PUT /api/profile", "Authenticated", "Xem/cập nhật profile."),
            ("POST /api/profile/kyc", "Volunteer", "Gửi KYC."),
        ]),
        ("Event", [
            ("GET /api/events", "Public", "Danh sách sự kiện."),
            ("POST /api/events", "Organizer", "Tạo event."),
            ("PUT /api/events/{id}", "Organizer", "Sửa event."),
            ("PUT /api/events/{id}/approve", "Admin", "Duyệt event."),
            ("PUT /api/events/{id}/complete", "Organizer/Admin", "Hoàn thành event."),
        ]),
        ("Registration/Attendance", [
            ("POST /api/events/{id}/register", "Volunteer", "Đăng ký."),
            ("PUT /api/events/{id}/registrations/{regId}/confirm", "Organizer", "Xác nhận."),
            ("POST /api/events/{id}/self-checkin", "Volunteer", "Tự check-in."),
            ("POST /api/events/{id}/registrations/{regId}/checkin", "Organizer", "Check-in."),
            ("POST /api/events/{id}/registrations/{regId}/checkout", "Organizer", "Check-out."),
        ]),
        ("Finance", [
            ("POST /api/events/{id}/support-campaigns", "Organizer", "Tạo campaign."),
            ("POST /api/support-campaigns/{id}/donations", "Volunteer", "Ủng hộ."),
            ("PUT /api/donations/{id}/confirm", "Organizer", "Xác nhận donation."),
            ("POST /api/events/{id}/sponsorship-proposals/sponsor-offer", "Sponsor", "Sponsor offer."),
            ("PUT /api/sponsorship-proposals/{id}/received", "Organizer", "Ghi nhận đã nhận."),
        ]),
        ("Admin", [
            ("GET /api/admin/users", "Admin", "Danh sách user."),
            ("PUT /api/admin/users/{id}/toggle-status", "Admin", "Khóa/mở user."),
            ("GET /api/admin/monitoring/summary", "Admin", "Monitoring."),
            ("GET /api/admin/audit-logs", "Admin", "Audit logs."),
            ("GET /api/admin/export/finance", "Admin", "Export tài chính."),
        ]),
    ]
    for group, rows in api_groups:
        add_page_section(doc, group, level=2)
        add_table(doc, ["Endpoint", "Role", "Mục đích"], rows, [4200, 1800, 3360])
        doc.add_heading("Xử lý lỗi chung", level=3)
        add_bullets(
            doc,
            [
                "401 khi chưa đăng nhập, token hết hạn hoặc user inactive.",
                "403 khi role không có quyền.",
                "400 khi dữ liệu đầu vào hoặc trạng thái nghiệp vụ không hợp lệ.",
                "404 khi tài nguyên không tồn tại hoặc không thuộc phạm vi actor.",
            ],
        )

    doc.add_page_break()
    doc.add_heading("ĐẶC TẢ MÀN HÌNH CHI TIẾT", level=1)
    screens = [
        ("LandingPage", "Guest/Public", ["Giới thiệu hệ thống", "CTA đăng ký theo role", "Sự kiện nổi bật", "Tác động cộng đồng"]),
        ("EventList", "Guest/Authenticated", ["Filter keyword/category/skill/location", "Danh sách card", "Bản đồ Leaflet", "Recommended events cho volunteer"]),
        ("EventDetail", "Guest/Volunteer", ["Thông tin event", "Organizer", "Impact", "Đăng ký/hủy đăng ký", "Self check-in khi đủ điều kiện"]),
        ("Profile", "Volunteer", ["Thông tin cá nhân", "Kỹ năng", "KYC", "Volunteer Passport"]),
        ("Activity", "Volunteer", ["My registrations", "Cancel request", "My donations", "Rating sau event"]),
        ("Achievements", "Volunteer", ["Badges", "Certificates", "Download/verify certificate"]),
        ("MyEvents", "Organizer", ["Danh sách event của tôi", "Trạng thái", "Quick actions", "Đi vào quản lý"]),
        ("ManageEvent", "Organizer", ["Registrations", "Check-in", "Shifts", "Campaigns", "Corporate", "Report"]),
        ("MySponsorships", "Sponsor", ["Proposal của tôi", "Accept/reject response", "Tracking trạng thái"]),
        ("AdminEvents", "Admin", ["Duyệt/từ chối", "Cancel/complete/uncomplete", "Transfer event", "Overdue preview"]),
        ("AdminVerifications", "Admin", ["Organizer verification", "Volunteer KYC", "Skill verification", "Request changes"]),
        ("AdminFinanceWatch", "Admin", ["Stale donations", "Unreported campaigns", "Open proposals past event"]),
        ("AdminMonitoring", "Admin", ["Health", "Summary", "Audit logs"]),
    ]
    for screen, role, features in screens:
        add_page_section(doc, screen, level=2)
        add_table(doc, ["Mục", "Nội dung"], [("Role", role), ("Component", screen), ("Nguồn", "BaseCore.WebClient/src/pages hoặc components liên quan")], [1800, 7560])
        doc.add_heading("Thành phần giao diện", level=3)
        add_bullets(doc, features)
        doc.add_heading("Trạng thái UI cần xử lý", level=3)
        add_bullets(doc, ["Loading", "Empty state", "Validation error", "Permission denied/redirect", "Success notification/toast"])

    doc.add_page_break()
    doc.add_heading("QUYẾT ĐỊNH THIẾT KẾ VÀ TRACEABILITY", level=1)
    doc.add_paragraph(
        "Phần này làm rõ các quyết định thiết kế then chốt theo cách tiếp cận architectural design, detailed design "
        "và UML design. Mỗi quyết định được gắn với lý do, trade-off và điểm kiểm thử để tránh tài liệu chỉ dừng ở mô tả tĩnh."
    )
    add_table(
        doc,
        ["Mã", "Quyết định", "Lý do", "Rủi ro/kiểm soát"],
        [
            ("ADR-01", "Dùng React/Vite làm frontend SPA và gọi API qua /api.", "Giảm coupling với port service, dễ đổi gateway khi deploy.", "Interceptor phải xử lý 401/refresh, route guard không thay thế backend authorization."),
            ("ADR-02", "Dùng Ocelot Gateway làm entrypoint cho service.", "Frontend có một API surface thống nhất, phù hợp mô hình client-server/layered.", "Route fallback phải đặt sau route chuyên biệt để không nuốt endpoint mới."),
            ("ADR-03", "Tách Auth/Event/Finance theo business capability.", "Theo dõi module rõ rệt, mỗi service đóng gói controller/service của miền nghiệp vụ.", "Dùng shared entity/common cần giữ contract ổn định, tránh sửa entity gây vỡ nhiều service."),
            ("ADR-04", "Dùng chung database trong phạm vi đồ án.", "Giảm độ phức tạp distributed transaction và phù hợp quy mô demo.", "Khi mở rộng cần xem xét bounded context/database per service."),
            ("ADR-05", "Controller mỏng, service chứa business logic.", "Dễ test luồng nghiệp vụ và đồng bộ quy tắc giữa API/UI.", "Không để controller bỏ qua service khi update state quan trọng."),
            ("ADR-06", "State transition được validate ở backend.", "UML state machine cho thấy trạng thái là ràng buộc nghiệp vụ, không phải chỉ là label UI.", "Mỗi transition phải có guard, side effect, notification/audit và test case."),
            ("ADR-07", "Certificate verify là endpoint public.", "Chứng chỉ cần được bên thứ ba kiểm chứng mà không cần tài khoản.", "Chỉ trả dữ liệu cần thiết; không rò rỉ thông tin nhạy cảm của volunteer."),
            ("ADR-08", "Audit log cho thao tác vận hành/nhạy cảm.", "Đáp ứng NFR auditability và minh bạch của hệ thống tình nguyện.", "Metadata cần gọn, không lưu password/token/proof file raw trong log."),
        ],
        [900, 2600, 3000, 2860],
    )

    doc.add_heading("Hợp đồng service và DTO", level=2)
    add_table(
        doc,
        ["Service", "Request/Command", "Entity chịu tác động", "Response/Side effect"],
        [
            ("AuthService", "Login/Register/RefreshToken request", "Users, RefreshTokens", "JWT pair, role claims, audit login failure khi cần."),
            ("ProfileService", "Update profile, SubmitKyc, SubmitSkillEvidence", "VolunteerProfiles, VolunteerSkills", "Notification cho admin/volunteer khi review."),
            ("OrganizerVerificationService", "SubmitVerification, ReviewVerification", "OrganizerVerifications", "Chuyển status, lưu reviewer, gửi notification kết quả."),
            ("EventService", "Create/Update/Approve/Reject/Complete/Cancel", "Events, Channels, Registrations, Certificates", "Sinh QR/channel khi approve; cấp/thu hồi certificate khi complete/uncomplete."),
            ("RegistrationService", "Register/Confirm/Cancel/CheckIn/CheckOut", "Registrations, WorkShifts, VolunteerProfiles", "Cập nhật giờ, metadata điểm danh, notification và audit."),
            ("FinanceService", "CreateCampaign, Donate, ConfirmDonation, CreateProposal, MarkReceived", "Campaigns, Donations, Proposals", "Impact amount chỉ cập nhật theo trạng thái hợp lệ."),
            ("ChannelService", "CreatePost, Comment, Like, PollVote", "Channels, Posts, Comments, Likes, Polls", "Broadcast SignalR cho thành viên có quyền truy cập channel."),
            ("AdminService", "ToggleUser, ReviewEntity, Export, Monitor", "Users, AuditLogs, domain entities", "Không xóa lịch sử, ghi audit, trả dữ liệu export có giới hạn."),
        ],
        [1900, 2600, 2500, 2360],
    )

    doc.add_heading("Guard condition và side effect theo state machine", level=2)
    add_table(
        doc,
        ["Đối tượng", "Transition", "Guard condition", "Side effect bắt buộc"],
        [
            ("Event", "Pending -> Approved", "Admin, organizer active/verified, dữ liệu thời gian/địa điểm hợp lệ.", "Sinh QR, tạo channel, public listing, notify organizer."),
            ("Event", "Approved -> Completed", "Event đã kết thúc hoặc admin/organizer xác nhận đủ điều kiện.", "Cấp certificate cho attended, cập nhật badge/passport, đóng các flow liên quan."),
            ("Event", "Completed -> Approved", "Admin rollback khi phát hiện lỗi hoàn thành.", "Thu hồi certificate đã cấp, ghi audit rõ lý do."),
            ("Registration", "Pending -> Confirmed", "Organizer sở hữu event, capacity/shift còn chỗ, volunteer không bị khóa.", "Notify volunteer, cho phép check-in trong cửa sổ hợp lệ."),
            ("Registration", "Confirmed -> Attended/CheckedOut", "QR/GPS/time window hợp lệ hoặc organizer manual override có lý do.", "Ghi metadata check-in/out, tính VolunteerHours, audit."),
            ("Campaign", "Open -> Closed/Reported", "Organizer sở hữu event, số tiền sử dụng không vượt số tiền confirmed.", "Public report, khóa donation mới nếu closed."),
            ("Donation", "PendingConfirmation -> Confirmed", "Organizer đối soát proof, amount > 0, campaign còn hợp lệ.", "Cộng vào impact, notify donor, audit finance."),
            ("Proposal", "Accepted -> Received -> Reported", "Organizer xác nhận đã nhận tài trợ và sau đó báo cáo sử dụng.", "Cộng corporate impact, cập nhật admin finance watch."),
        ],
        [1600, 2200, 3200, 2360],
    )

    doc.add_heading("Liên hệ giữa thiết kế và kiểm thử", level=2)
    doc.add_paragraph(
        "Khi chuyển từ thiết kế sang kiểm thử, nhóm tập trung vào những luồng có rủi ro nghiệp vụ cao: đăng nhập, đăng ký sự kiện, "
        "điểm danh, hoàn thành sự kiện, tài trợ, xác minh và phân quyền. Các luồng này được ưu tiên vì nếu sai sẽ ảnh hưởng trực tiếp "
        "đến uy tín của hệ thống hoặc làm sai lệch kết quả đóng góp."
    )
    add_table(
        doc,
        ["Luồng chính", "Thiết kế liên quan", "Cách kiểm thử dự kiến", "Kết quả mong đợi"],
        [
            ("Đăng nhập và phân quyền", "JWT middleware, refresh token, ProtectedRoute, IsActive middleware.", "/api/auth/login, /api/auth/refresh, Login/Register UI", "Đúng role vào đúng khu vực; user bị khóa không dùng được phiên cũ."),
            ("Đăng ký sự kiện", "RegistrationService kiểm tra trạng thái event, KYC, capacity, shift và kỹ năng.", "EventDetail, POST /events/{id}/register", "Đăng ký hợp lệ thành Pending; trường hợp full hoặc thiếu điều kiện trả lỗi rõ."),
            ("Điểm danh", "QR/GPS/time-window guard, metadata điểm danh và audit.", "Self check-in UI, organizer check-in tab", "Chỉ volunteer đã Confirmed mới check-in; QR sai hoặc ngoài bán kính bị chặn."),
            ("Hoàn thành sự kiện", "EventService phối hợp CertificateService và BadgeService.", "Admin/Organizer complete action, certificate verify", "Người tham gia hợp lệ có chứng chỉ; rollback sẽ thu hồi chứng chỉ đã cấp."),
            ("Donation", "Campaign/Donation state machine và tính toán impact.", "Campaign UI, donation confirm API", "Chỉ donation Confirmed xuất hiện trong impact công khai."),
            ("Tài trợ doanh nghiệp", "Proposal state machine, quyền theo vai trò và admin finance watch.", "MySponsorships, ManageEvent Corporate tab", "Received/Reported mới tính vào impact; duplicate proposal bị chặn."),
            ("Xác minh", "Các luồng approve/reject/request changes có reason và notification.", "AdminVerifications, profile/organizer verification UI", "Trạng thái đổi đúng và người liên quan nhận thông báo."),
            ("API nhạy cảm", "Role policy, backend ownership check, audit và rate limit.", "Admin/user/event/finance protected APIs", "Sai role bị từ chối; thao tác nhạy cảm có dấu vết kiểm tra."),
        ],
        [1600, 3000, 2500, 2260],
    )

    doc.add_heading("Lưu ý khi đọc các sơ đồ thiết kế", level=2)
    doc.add_paragraph(
        "Các sơ đồ trong tài liệu được dùng để giải thích cách hệ thống vận hành ở mức thiết kế, không thay thế hoàn toàn cho mã nguồn. "
        "Vì vậy khi đọc sơ đồ cần phân biệt rõ vai trò người dùng, đối tượng dữ liệu, luồng tương tác và trạng thái nghiệp vụ."
    )
    add_bullets(
        doc,
        [
            "Actor không được nhầm với entity: Guest, Volunteer, Organizer, Sponsor, Admin là actor; User/Event/Registration là class/entity.",
            "Use case là mục tiêu của actor, không phải tên nút UI hay tên hàm controller.",
            "Activity flow phải thể hiện decision node cho lỗi nghiệp vụ chính: full capacity, thiếu KYC, sai QR/GPS, finance chưa xác nhận.",
            "Sequence flow phải có đủ boundary UI, controller/API, service, database và notification/audit nếu có side effect.",
            "Class/entity diagram phải thể hiện cardinality quan trọng: Event-Registration, User-VolunteerProfile, Campaign-Donation, Proposal-Sponsor.",
            "State machine phải có trạng thái cuối hoặc trạng thái khóa nghiệp vụ như Completed, Cancelled, Reported, Rejected.",
            "Các yêu cầu chất lượng như bảo mật, hiệu năng và minh bạch được thể hiện qua kiến trúc, ràng buộc dữ liệu và kịch bản kiểm thử.",
            "Database design phải phân biệt unique constraint, foreign key và rule validate ở service.",
        ],
    )

def build_srs() -> Document:
    doc = setup_document("ĐẶC TẢ YÊU CẦU PHẦN MỀM", "Software Requirements Specification - SRS")
    add_history(doc)
    add_signature_page(doc)
    add_toc_note(doc)
    intro_common(doc, "đặc tả yêu cầu phần mềm")
    add_natural_project_context(doc, "SRS")

    doc.add_heading("TỔNG QUAN HỆ THỐNG", level=1)
    doc.add_heading("Bối cảnh", level=2)
    doc.add_paragraph(
        "Các hoạt động tình nguyện thường được tổ chức rời rạc qua mạng xã hội, khó kiểm chứng uy tín tổ chức, "
        "khó ghép đúng tình nguyện viên theo kỹ năng và thiếu chứng cứ chính thống về đóng góp. VolunteerHub giải quyết "
        "bài toán này bằng một cổng web minh bạch kết nối người tham gia, nhà tổ chức, nhà tài trợ và quản trị viên."
    )
    doc.add_heading("Mục tiêu", level=2)
    add_numbered(
        doc,
        [
            "Kết nối đúng người, đúng việc theo kỹ năng, vị trí và nhu cầu sự kiện.",
            "Minh bạch hóa thông tin tổ chức, sự kiện, tài trợ, ủng hộ và kết quả tác động.",
            "Quản lý trọn vòng đời sự kiện từ tạo, duyệt, đăng ký, điểm danh, check-out đến hoàn thành.",
            "Số hóa đóng góp bằng giờ tình nguyện, chứng chỉ điện tử, huy hiệu và hộ chiếu tình nguyện.",
            "Hỗ trợ vận hành tại hiện trường bằng giao diện responsive, QR/GPS check-in và thông báo.",
        ],
    )
    doc.add_heading("Tác nhân hệ thống", level=2)
    actor_rows = [
        ("Guest", "Xem landing page, sự kiện công khai, impact, xác thực chứng chỉ."),
        ("Volunteer", "Quản lý hồ sơ/kỹ năng/KYC, tìm và đăng ký sự kiện, điểm danh, xem passport, chứng chỉ, huy hiệu, ủng hộ."),
        ("Organizer", "Xác minh tổ chức, tạo sự kiện, quản lý đăng ký/ca/điểm danh, tạo campaign, xử lý tài trợ, báo cáo."),
        ("Sponsor", "Quản lý hồ sơ nhà tài trợ, gửi/nhận đề nghị tài trợ, theo dõi tài trợ."),
        ("Admin", "Duyệt/từ chối dữ liệu, quản trị user/danh mục/huy hiệu/rating, giám sát tài chính, monitoring, export."),
    ]
    add_table(doc, ["Actor", "Vai trò"], actor_rows, [1800, 7560])
    doc.add_heading("Phạm vi", level=2)
    doc.add_heading("Trong phạm vi", level=3)
    add_bullets(
        doc,
        [
            "Đăng ký, đăng nhập, refresh token và phân quyền theo vai trò.",
            "Quản lý hồ sơ tình nguyện viên, kỹ năng, KYC, xác minh tổ chức và hồ sơ sponsor.",
            "Quản lý sự kiện, đăng ký, ca làm việc, điểm danh QR/GPS, check-out, hoàn thành và chứng chỉ.",
            "Quản lý channel trao đổi, thông báo, dashboard, rating hai chiều, huy hiệu.",
            "Quản lý campaign ủng hộ cá nhân, donation, sponsorship proposal và giám sát tài chính.",
            "Quản trị người dùng, danh mục, kỹ năng, audit log, monitoring, export JSON/CSV.",
        ],
    )
    doc.add_heading("Ngoài phạm vi", level=3)
    add_bullets(
        doc,
        [
            "Thanh toán online thật; hệ thống chỉ ghi nhận khoản đã chuyển/đã nhận.",
            "Ứng dụng mobile native và push notification native.",
            "Hợp đồng tài trợ pháp lý đầy đủ.",
            "Workflow khiếu nại/ticket phức tạp ngoài các thao tác moderation hiện có.",
        ],
    )

    doc.add_heading("YÊU CẦU CHỨC NĂNG", level=1)
    fr_rows = [
        ("FR-01", "Xác thực và phân quyền", "Đăng ký, đăng nhập, JWT access token, refresh token, role guard, chặn user bị khóa."),
        ("FR-02", "Trang công khai và khám phá sự kiện", "Landing page, danh sách Approved/Completed, filter, bản đồ, chi tiết, chia sẻ link."),
        ("FR-03", "Hồ sơ tình nguyện viên", "Cập nhật thông tin, avatar, kỹ năng, minh chứng kỹ năng, trạng thái xác minh."),
        ("FR-04", "Volunteer Passport", "Lịch sử tham gia, tổng giờ, chứng chỉ, huy hiệu."),
        ("FR-05", "KYC volunteer", "Gửi CCCD/chân dung, admin approve/reject/request changes, notify."),
        ("FR-06", "Organizer verification", "Gửi hồ sơ tổ chức, admin duyệt/từ chối/yêu cầu bổ sung, chỉ Verified mới được duyệt event."),
        ("FR-07", "Quản lý sự kiện", "Organizer tạo/sửa/xóa có điều kiện, resubmit, admin approve/reject/cancel/transfer/complete."),
        ("FR-08", "Đăng ký sự kiện", "Volunteer đăng ký Approved event, kiểm tra thời gian, KYC, capacity, shift và kỹ năng."),
        ("FR-09", "Quản lý đăng ký", "Organizer xác nhận, hủy, xử lý yêu cầu hủy, chuyển ca, walk-in."),
        ("FR-10", "Ca làm việc", "Tạo/sửa/xóa ca trước khi có đăng ký, bắt buộc chọn ca nếu event đã chia ca."),
        ("FR-11", "Điểm danh và check-out", "QR/GPS/self check-in, manual attend, check-out tính giờ thực tế, audit metadata."),
        ("FR-12", "Hoàn thành sự kiện", "Complete Approved event, bỏ qua pending, cấp chứng chỉ cho attended, rollback bởi admin."),
        ("FR-13", "Chứng chỉ và huy hiệu", "Cấp PDF certificate, verify code, badge tự động theo số sự kiện/giờ."),
        ("FR-14", "Rating hai chiều", "Volunteer đánh giá organizer, organizer đánh giá volunteer, admin hide/unhide/delete."),
        ("FR-15", "Channel trao đổi", "Channel tự tạo khi event Approved, post/comment/like/poll, SignalR realtime."),
        ("FR-16", "Donation campaign", "Organizer tạo campaign, volunteer ủng hộ, organizer confirm/reject, report sử dụng."),
        ("FR-17", "Sponsorship proposal", "Organizer request hoặc Sponsor offer, accept/reject/received/report/cancel."),
        ("FR-18", "Dashboard và thông báo", "Dashboard theo role, notification list, mark read, điều hướng theo loại thông báo."),
        ("FR-19", "Admin vận hành", "User, catalog, skill, badge, verification, finance watch, monitoring, audit log, export."),
        ("FR-20", "Upload", "Upload ảnh/file cho event cover, KYC, evidence, donation proof."),
        ("FR-21", "Gợi ý", "Gợi ý event theo skill và knowledge graph khi khả dụng."),
    ]
    add_table(doc, ["Mã", "Tên yêu cầu", "Mô tả"], fr_rows, [950, 2300, 6110])

    doc.add_heading("MÔ HÌNH USE CASE", level=1)
    add_figure(
        doc,
        DIAGRAM_DIR / "uml_use_case.png",
        "Hình 1. Use Case Diagram tổng quan theo actor và system boundary của VolunteerHub.",
    )
    uc_rows = [
        ("UC-01", "Guest", "Xem danh sách và chi tiết sự kiện công khai"),
        ("UC-02", "Guest", "Xác thực chứng chỉ bằng mã verify"),
        ("UC-03", "Volunteer", "Đăng ký/đăng nhập và cập nhật hồ sơ"),
        ("UC-04", "Volunteer", "Gửi KYC và minh chứng kỹ năng"),
        ("UC-05", "Volunteer", "Tìm kiếm, đăng ký, rút hoặc xin hủy đăng ký sự kiện"),
        ("UC-06", "Volunteer", "Tự check-in và xem hoạt động/chứng chỉ/huy hiệu"),
        ("UC-07", "Volunteer", "Ủng hộ campaign và theo dõi donation"),
        ("UC-08", "Organizer", "Gửi xác minh tổ chức"),
        ("UC-09", "Organizer", "Tạo, sửa và gửi duyệt sự kiện"),
        ("UC-10", "Organizer", "Quản lý đăng ký, ca, điểm danh, check-out"),
        ("UC-11", "Organizer", "Hoàn thành sự kiện và xem báo cáo tác động"),
        ("UC-12", "Organizer", "Tạo campaign và xử lý donation"),
        ("UC-13", "Organizer/Sponsor", "Xử lý sponsorship proposal"),
        ("UC-14", "Admin", "Duyệt event, KYC, skill, organizer verification"),
        ("UC-15", "Admin", "Quản trị user, catalog, badge, rating, finance, monitoring, export"),
    ]
    add_table(doc, ["Mã UC", "Actor", "Use case"], uc_rows, [1100, 1800, 6460])

    doc.add_heading("Đặc tả Use Case tiêu biểu", level=2)
    doc.add_heading("Activity diagram nghiệp vụ chính", level=2)
    add_figure(
        doc,
        DIAGRAM_DIR / "uml_activity_registration.png",
        "Hình 2. Activity Diagram cho luồng đăng ký, xác nhận và điểm danh sự kiện.",
    )

    use_cases = [
        (
            "UC-09 - Tạo và gửi duyệt sự kiện",
            [
                ("Actor chính", "Organizer"),
                ("Tiền điều kiện", "Organizer đã đăng nhập, tài khoản active, hồ sơ tổ chức đủ điều kiện vận hành."),
                ("Luồng chính", "Organizer nhập thông tin sự kiện; hệ thống validate; lưu Event ở trạng thái Pending; admin xem trong hàng chờ duyệt."),
                ("Luồng thay thế", "Nếu event bị Rejected, organizer sửa và resubmit về Pending."),
                ("Ngoại lệ", "Thiếu dữ liệu bắt buộc, thời gian không hợp lệ, organizer không active."),
                ("Hậu điều kiện", "Sự kiện được lưu để chờ admin phê duyệt."),
            ],
        ),
        (
            "UC-10 - Đăng ký và xác nhận tham gia",
            [
                ("Actor chính", "Volunteer, Organizer"),
                ("Tiền điều kiện", "Event Approved, chưa bắt đầu, chưa hết chỗ, volunteer thỏa KYC/shift/skill nếu có."),
                ("Luồng chính", "Volunteer gửi đăng ký Pending; organizer xem danh sách; organizer Confirm; hệ thống thông báo volunteer."),
                ("Luồng thay thế", "Volunteer rút khi Pending; sau Confirm phải gửi yêu cầu hủy để organizer xử lý."),
                ("Ngoại lệ", "Event full, sai shift, chưa KYC, đăng ký trùng chưa Cancelled."),
                ("Hậu điều kiện", "Registration chuyển Confirmed hoặc Cancelled theo xử lý."),
            ],
        ),
        (
            "UC-11 - Điểm danh và check-out",
            [
                ("Actor chính", "Volunteer, Organizer"),
                ("Tiền điều kiện", "Registration Confirmed, event Approved, nằm trong cửa sổ check-in."),
                ("Luồng chính", "Người dùng quét QR hoặc cung cấp GPS; hệ thống validate; ghi AttendedAt; check-out ghi CheckedOutAt và tính VolunteerHours."),
                ("Luồng thay thế", "Organizer walk-in volunteer hoặc manual attend trong 7 ngày sau event."),
                ("Ngoại lệ", "QR sai, ngoài bán kính, ngoài cửa sổ thời gian, chưa Confirmed, đã check-in/check-out."),
                ("Hậu điều kiện", "Giờ tình nguyện được ghi nhận và dùng cho passport/badge/certificate."),
            ],
        ),
        (
            "UC-13 - Tài trợ doanh nghiệp",
            [
                ("Actor chính", "Organizer, Sponsor"),
                ("Tiền điều kiện", "Event Approved; sponsor active; chưa có proposal active trùng."),
                ("Luồng chính", "Organizer gửi lời mời hoặc Sponsor gửi đề nghị; bên nhận accept/reject; organizer đánh dấu Received; sau event báo cáo sử dụng."),
                ("Luồng thay thế", "Bên tạo cancel khi Pending; admin có thể revert Accepted về Pending."),
                ("Ngoại lệ", "Sponsor hủy sau Accepted, số tiền không hợp lệ, proposal duplicate."),
                ("Hậu điều kiện", "Chỉ proposal Received/Reported được tính vào impact public."),
            ],
        ),
    ]
    for heading, rows in use_cases:
        doc.add_heading(heading, level=3)
        add_table(doc, ["Mục", "Nội dung"], rows, [2100, 7260])

    doc.add_heading("TRẠNG THÁI VÀ QUY TẮC NGHIỆP VỤ", level=1)
    state_rows = [
        ("Event", "Pending -> Approved -> Completed; Pending -> Rejected; Approved/Pending -> Cancelled."),
        ("Registration", "Pending -> Confirmed -> attended/check-out; Pending/Confirmed -> Cancelled; Confirmed có thể CancelRequested."),
        ("SupportCampaign", "Draft -> Open -> Closed -> Reported; Draft/Open -> Cancelled."),
        ("IndividualDonation", "PendingConfirmation -> Confirmed/Rejected/Cancelled."),
        ("SponsorshipProposal", "Pending -> Accepted -> Received -> Reported; Pending -> Rejected/Cancelled; Accepted -> Cancelled khi event hủy."),
        ("Verification", "Unverified/SelfDeclared -> PendingVerification -> Verified/Rejected/ChangesRequested."),
    ]
    add_table(doc, ["Đối tượng", "State machine"], state_rows, [2100, 7260])
    doc.add_heading("Quy tắc chính", level=2)
    add_numbered(
        doc,
        [
            "Chỉ sự kiện Approved mới mở đăng ký, điểm danh và tài trợ.",
            "Organizer chỉ được quản lý sự kiện do mình sở hữu; admin có quyền vận hành cấp hệ thống.",
            "Registration chỉ được check-in khi Confirmed và nằm trong cửa sổ hợp lệ.",
            "Check-out mới tính giờ thực tế; check-in không cộng giờ ngay.",
            "Certificate chỉ cấp cho registration đã attended và có giờ hợp lệ.",
            "Donation chỉ tính public khi Confirmed; proposal chỉ tính public khi Received/Reported.",
            "Dữ liệu đã phát sinh lịch sử ưu tiên hủy/ẩn/chặn xóa thay vì xóa cứng.",
            "Reject/request changes phải có lý do rõ ràng, tối thiểu 10 ký tự với các flow nhạy cảm.",
        ],
    )

    doc.add_heading("YÊU CẦU DỮ LIỆU", level=1)
    data_rows = [
        ("User", "Tài khoản, vai trò, trạng thái active, thông tin liên hệ."),
        ("VolunteerProfile", "Hồ sơ volunteer, avatar, nhóm máu, ngôn ngữ, KYC, tổng giờ."),
        ("VolunteerSkill/Skill", "Danh mục kỹ năng và kỹ năng người dùng, trạng thái xác minh."),
        ("OrganizerVerification", "Hồ sơ xác minh pháp lý/tổ chức của organizer."),
        ("Event/EventCategory", "Sự kiện, trạng thái, địa điểm, tọa độ, kỹ năng yêu cầu, QR, organizer."),
        ("WorkShift", "Ca làm việc thuộc event, thời gian, capacity, skill yêu cầu."),
        ("Registration", "Đăng ký, trạng thái, shift, check-in/out, giờ, yêu cầu hủy."),
        ("Certificate/Badge/UserBadge", "Chứng chỉ, mã verify, huy hiệu và lịch sử trao."),
        ("Channel/Post/Comment/Like/Poll", "Trao đổi trong channel sự kiện."),
        ("SupportCampaign/IndividualDonation", "Đợt kêu gọi và khoản ủng hộ cá nhân."),
        ("SponsorProfile/SponsorshipProposal", "Hồ sơ sponsor và đề nghị tài trợ doanh nghiệp."),
        ("Rating/Notification/AuditLog", "Đánh giá, thông báo, nhật ký thao tác."),
    ]
    add_table(doc, ["Thực thể", "Ý nghĩa"], data_rows, [2600, 6760])

    doc.add_heading("YÊU CẦU GIAO DIỆN", level=1)
    add_bullets(
        doc,
        [
            "Giao diện public gồm landing, danh sách sự kiện, bản đồ, chi tiết event và xác thực chứng chỉ.",
            "Volunteer có dashboard, profile, activity, achievements và màn hình đăng ký/check-in.",
            "Organizer có my-events, event form, manage event với các tab đăng ký, check-in, campaign, corporate, report.",
            "Sponsor có sponsor profile và my sponsorships.",
            "Admin có các màn events, users, verifications, catalog, ratings, finance, monitoring, export.",
            "Giao diện phải responsive, ưu tiên thao tác tại hiện trường trên mobile.",
        ],
    )

    doc.add_heading("YÊU CẦU PHI CHỨC NĂNG", level=1)
    nfr_rows = [
        ("Bảo mật", "JWT, refresh token, role-based access, không trả password/salt, chặn user inactive."),
        ("Minh bạch", "Impact public, donation/proposal chỉ tính tiền đã xác nhận, certificate verify độc lập."),
        ("Hiệu năng", "Pagination/filter, giới hạn export, rate limiting cho API nhạy cảm."),
        ("Khả dụng", "Responsive desktop/mobile, thao tác QR/GPS thuận tiện tại hiện trường."),
        ("Bảo trì", "Tách service, API contract ổn định, tài liệu được cập nhật theo luồng nghiệp vụ."),
        ("Kiểm toán", "Audit log cho thao tác nhạy cảm, check-in metadata gồm method/GPS/IP khi có."),
    ]
    add_table(doc, ["Nhóm yêu cầu", "Mô tả"], nfr_rows, [2200, 7160])

    doc.add_heading("TIÊU CHÍ NGHIỆM THU", level=1)
    add_bullets(
        doc,
        [
            "Guest xem được event Approved và verify certificate.",
            "Volunteer đăng ký được event hợp lệ, bị chặn đúng khi thiếu KYC hoặc event full.",
            "Organizer xác nhận, điểm danh, check-out và hoàn thành event đúng state machine.",
            "Certificate được cấp/thu hồi đúng khi complete/uncomplete.",
            "Campaign và proposal phản ánh đúng số tiền Confirmed/Received trong impact.",
            "Admin duyệt/từ chối/yêu cầu bổ sung đúng các luồng verification và export được dữ liệu.",
        ],
    )
    add_srs_functional_detail(doc)
    add_detailed_srs_appendix(doc)
    return doc


def build_detail_design() -> Document:
    doc = setup_document("THIẾT KẾ CHI TIẾT", "Detail Design Specification")
    add_history(doc)
    add_signature_page(doc)
    add_toc_note(doc)
    intro_common(doc, "thiết kế chi tiết")
    add_natural_project_context(doc, "Detail Design")

    doc.add_heading("TỔNG QUAN HỆ THỐNG", level=1)
    doc.add_paragraph(
        "VolunteerHub được triển khai theo kiến trúc gateway/service-ready. Frontend React/Vite gọi API qua đường dẫn /api, "
        "Vite proxy về API Gateway Ocelot. Gateway định tuyến theo miền sang Identity, Event, Finance hoặc Legacy API; "
        "source controller gốc đặt trong BaseCore.APIService/Controllers và được link sang các service bằng Compile Include trong .csproj. "
        "Các service dùng chung SQL Server database VolunteerHub thông qua EF Core DbContext và chia sẻ các project Entities, Repository, Services, Common."
    )
    add_table(
        doc,
        ["Thành phần", "Project", "Port", "Trách nhiệm"],
        [
            ("Frontend", "BaseCore.WebClient", "3000", "UI chính, route theo role, gọi API qua axios."),
            ("Gateway", "BaseCore.ApiGateway", "5000", "Routing Ocelot, CORS, entrypoint /api."),
            ("Identity", "BaseCore.AuthService", "5002", "Auth, profile, KYC, organizer verification, user, notification, monitoring."),
            ("Event", "BaseCore.EventService", "5003", "Event, registration, attendance, certificate, channel SignalR, dashboard, rating."),
            ("Finance", "BaseCore.FinanceService", "5004", "Campaign, donation, sponsorship proposal, sponsor profile, finance export."),
            ("Legacy", "BaseCore.APIService", "5001", "Fallback legacy và controller được link/shared cho service split."),
            ("Database", "SQL Server", "-", "Database VolunteerHub, EF Core migrations, seed demo."),
        ],
        [1600, 2300, 900, 4560],
    )

    doc.add_heading("KIẾN TRÚC HỆ THỐNG", level=1)
    doc.add_heading("Mô hình kiến trúc", level=2)
    add_bullets(
        doc,
        [
            "Client React/Vite dùng axios instance với baseURL /api và interceptor gắn JWT.",
            "API Gateway Ocelot định tuyến theo priority: auth/profile/user/skill về Identity; event/channel/dashboard về Event; finance về Finance.",
            "Mỗi service tự cấu hình CORS, JWT Bearer, rate limiting, DbContext và middleware IsActive.",
            "EventService map SignalR hub /hubs/channel cho realtime post/comment.",
            "Tất cả service dùng chung database để giảm độ phức tạp đồng bộ trong phạm vi đồ án.",
        ],
    )
    doc.add_heading("Gateway routing chính", level=2)
    add_table(
        doc,
        ["Nhóm route", "Service đích", "Ghi chú"],
        [
            ("/api/auth, /api/profile, /api/users, /api/skills", "Identity 5002", "Auth, profile, user, skill, verification."),
            ("/api/events, /api/event-categories, /api/my-registrations", "Event 5003", "Sự kiện, đăng ký, ca, attendance."),
            ("/api/certificates, /api/channels, /api/ratings, /api/dashboard", "Event 5003", "Engagement, certificate, dashboard."),
            ("/api/support-campaigns, /api/donations, /api/sponsorship-proposals", "Finance 5004", "Donation và tài trợ."),
            ("/api/admin/finance, /api/admin/export/finance", "Finance 5004", "Giám sát và export tài chính."),
            ("/api/admin/export/events", "Legacy 5001", "Export event fallback."),
            ("/api/{everything}", "Legacy 5001", "Fallback route."),
        ],
        [3200, 1900, 4260],
    )

    doc.add_heading("THIẾT KẾ MODULE", level=1)
    add_figure(
        doc,
        DIAGRAM_DIR / "uml_component_deployment.png",
        "Hình 3. Component/Deployment Diagram cho frontend, gateway, service và database.",
    )

    module_rows = [
        ("Identity/Auth", "AuthController, UserController, ProfileController", "Đăng ký, đăng nhập, refresh token, profile, KYC, skill verification."),
        ("Organizer Verification", "OrganizerVerificationController", "Gửi hồ sơ tổ chức, admin review, request changes, notification."),
        ("Event Lifecycle", "EventsController, EventService", "Search, create, approve, reject, cancel, complete, uncomplete, transfer, QR rotate."),
        ("Registration/Attendance", "RegistrationsController, RegistrationService", "Register, withdraw, confirm, cancel, check-in, self-check-in, walk-in, checkout, manual attend."),
        ("WorkShift", "WorkShiftsController", "CRUD ca làm việc, kiểm tra event ownership và ràng buộc registration."),
        ("Certificate/Badge", "CertificatesController, CertificateService, BadgeService", "Cấp PDF, verify code, trao huy hiệu theo điều kiện."),
        ("Channel", "ChannelsController, ChannelService, ChannelHub", "Post, comment, like, poll, pin, realtime notification."),
        ("Finance", "SupportCampaignController, SponsorshipProposalController, SponsorProfileController", "Campaign, donation, proposal, sponsor profile."),
        ("Admin", "AdminController, MonitoringController, AdminFinanceController", "User, verification, catalog, rating, monitoring, audit, export."),
        ("Frontend", "BaseCore.WebClient/src/App.jsx, services/api.js", "Route theo role, protected route, API facade."),
    ]
    add_table(doc, ["Module", "Thành phần chính", "Chức năng"], module_rows, [1900, 3100, 4360])

    doc.add_heading("THIẾT KẾ LỚP", level=1)
    doc.add_paragraph(
        "Thiết kế backend theo hướng controller mỏng gọi service nghiệp vụ; service thao tác qua EF Core DbContext/repository và phát notification/audit khi cần."
    )
    add_figure(
        doc,
        DIAGRAM_DIR / "uml_class_domain.png",
        "Hình 4. Class/Domain Model thể hiện các entity và quan hệ lõi.",
    )
    class_rows = [
        ("EventService", "SearchAsync, CreateAsync, ApproveAsync, RejectAsync, CompleteAsync, CancelAsync, UncompleteAsync", "Điều phối vòng đời event, channel tự tạo, cascade finance, certificate."),
        ("RegistrationService", "RegisterAsync, ConfirmAsync, CheckInAsync, SelfCheckInAsync, CheckOutAsync, WalkInAsync", "Đảm bảo state registration, capacity, shift, KYC, QR/GPS, tính giờ."),
        ("CertificateService", "IssueCertificatesForEventAsync, GetByCodeAsync, BuildCertificatePdf", "Cấp chứng chỉ, tạo PDF và verify certificate."),
        ("BadgeService", "CheckAndAwardAsync, GetAllAsync, GetByUserAsync", "Trao huy hiệu tự động theo điều kiện."),
        ("ChannelService", "CanAccessChannelAsync, GetPostsAsync, CreatePostAsync, ToggleLikeAsync", "Quyền truy cập channel và tương tác realtime."),
        ("AuditLogService", "RecordAsync", "Ghi log thao tác nhạy cảm với metadata."),
        ("NotificationService", "SendAsync", "Tạo thông báo cho user và điều hướng frontend."),
    ]
    add_table(doc, ["Lớp", "Phương thức chính", "Trách nhiệm"], class_rows, [1900, 3600, 3860])

    doc.add_heading("Thiết kế thuật toán cho phương thức chính", level=2)
    doc.add_paragraph(
        "Các bảng dưới đây mô tả sâu hơn cách một số phương thức nghiệp vụ được thực hiện. "
        "Mỗi phương thức không chỉ có tên và trách nhiệm, mà còn có điều kiện vào, các bước xử lý chính, dữ liệu cập nhật và tình huống lỗi."
    )
    algorithm_specs = [
        (
            "RegistrationService.RegisterAsync(eventId, userId, shiftId)",
            [
                ("Mục đích", "Tạo đăng ký tham gia sự kiện cho volunteer khi sự kiện còn mở và volunteer thỏa điều kiện."),
                ("Dữ liệu vào", "eventId, userId lấy từ JWT, shiftId nếu đăng ký theo ca."),
                ("Dữ liệu cập nhật", "Registrations, notification cho organizer/volunteer nếu cần."),
                ("Ngoại lệ", "Event không tồn tại, chưa được duyệt, đã kết thúc, hết chỗ, thiếu KYC, sai shift, đã đăng ký trước đó."),
            ],
            """BEGIN
  mở transaction mức Serializable
  đọc event kèm registration hiện có bằng khóa cập nhật
  IF event không tồn tại OR status != Approved THEN báo lỗi
  IF hiện tại nằm ngoài thời gian cho phép đăng ký THEN báo lỗi
  IF volunteer đã có registration active THEN báo lỗi
  IF event yêu cầu KYC AND volunteer chưa Verified THEN báo lỗi
  IF shiftId có giá trị THEN
      kiểm shift thuộc event và còn chỗ
      kiểm skill yêu cầu của shift nếu có
  ELSE
      kiểm capacity tổng của event
  END IF
  tạo Registration với Status = Pending
  lưu database
  commit transaction
  trả về registration mới
END""",
        ),
        (
            "RegistrationService.CheckInAsync(eventId, regId, payload)",
            [
                ("Mục đích", "Xác nhận volunteer có mặt tại sự kiện bằng QR hoặc GPS."),
                ("Dữ liệu vào", "eventId, regId, qrCode hoặc latitude/longitude."),
                ("Dữ liệu cập nhật", "AttendedAt, check-in metadata, VolunteerHours tạm thời, audit/notification."),
                ("Ngoại lệ", "Registration chưa Confirmed, đã check-in, ngoài cửa sổ thời gian, QR sai, GPS ngoài bán kính."),
            ],
            """BEGIN
  đọc registration, event và shift liên quan
  IF registration không thuộc event THEN báo lỗi
  IF registration.Status != Confirmed THEN báo lỗi
  IF registration.AttendedAt đã có THEN trả về trạng thái hiện tại
  kiểm cửa sổ check-in theo event/shift
  IF payload có qrCode THEN
      so sánh qrCode với event.QrCode
      IF không khớp THEN báo lỗi
  ELSE
      kiểm latitude/longitude hợp lệ
      tính khoảng cách Haversine đến tọa độ event
      IF khoảng cách > CheckInRadiusKm THEN báo lỗi
  END IF
  cập nhật AttendedAt = thời điểm hiện tại
  ghi method, IP/GPS metadata nếu có
  lưu database và ghi audit
  trả về registration đã cập nhật
END""",
        ),
        (
            "EventService.CompleteAsync(eventId, actorId)",
            [
                ("Mục đích", "Hoàn thành sự kiện, khóa luồng vận hành và cấp chứng chỉ cho người tham gia hợp lệ."),
                ("Dữ liệu vào", "eventId, actorId của organizer/admin."),
                ("Dữ liệu cập nhật", "Events.Status, Certificates, UserBadges, VolunteerProfile.TotalVolunteerHours."),
                ("Ngoại lệ", "Không có quyền, event chưa Approved, event chưa đủ điều kiện hoàn thành."),
            ],
            """BEGIN
  đọc event cùng registrations đã Confirmed/Attended
  kiểm actor là organizer sở hữu event hoặc admin
  IF event.Status != Approved THEN báo lỗi
  cập nhật event.Status = Completed
  FOR từng registration đã Attended
      IF chưa có CheckedOutAt THEN tính giờ theo khung hợp lệ
      cập nhật VolunteerHours
      tạo certificate nếu user/event chưa có certificate
      cập nhật tổng giờ trong VolunteerProfile
      gọi BadgeService kiểm và trao huy hiệu mới
  END FOR
  lưu database
  gửi notification hoàn thành/certificate cho volunteer
  ghi audit hoàn thành sự kiện
END""",
        ),
        (
            "SupportCampaignController.ConfirmDonation(donationId)",
            [
                ("Mục đích", "Xác nhận khoản ủng hộ sau khi organizer đối soát minh chứng."),
                ("Dữ liệu vào", "donationId, actorId từ JWT."),
                ("Dữ liệu cập nhật", "IndividualDonations.Status, ConfirmedAt, ConfirmedBy, campaign impact."),
                ("Ngoại lệ", "Không có quyền, donation không ở trạng thái PendingConfirmation, campaign đã hủy."),
            ],
            """BEGIN
  đọc donation kèm campaign và event
  kiểm actor là organizer sở hữu event hoặc admin
  IF donation.Status != PendingConfirmation THEN báo lỗi
  IF campaign.Status không cho nhận tiền THEN báo lỗi
  cập nhật donation.Status = Confirmed
  cập nhật ConfirmedAt, ConfirmedBy
  tính lại confirmed amount của campaign
  lưu database
  gửi notification cho donor
  ghi audit finance
END""",
        ),
    ]
    for name, rows, pseudo in algorithm_specs:
        doc.add_heading(name, level=3)
        add_table(doc, ["Thành phần", "Mô tả"], rows, [1800, 7560])
        doc.add_paragraph("Giả mã xử lý:")
        add_pseudocode(doc, pseudo)

    doc.add_heading("THIẾT KẾ DỮ LIỆU", level=1)
    entity_rows = [
        ("Users", "Id, UserName unique, Password, Salt, Name, Email, Phone, UserType, IsActive", "Tài khoản và phân quyền."),
        ("VolunteerProfiles", "UserId unique, KycStatus, Identity images, Avatar, TotalVolunteerHours", "Hồ sơ/KYC volunteer."),
        ("VolunteerSkills", "UserId + SkillId unique, VerificationStatus, EvidenceUrl", "Kỹ năng và xác minh kỹ năng."),
        ("OrganizerVerifications", "OrganizerId unique, Status, SubmittedAt, VerifiedBy", "Xác minh tổ chức."),
        ("Events", "OrganizerId, CategoryId, Status, StartDate, EndDate, QrCode, CheckInRadiusKm", "Sự kiện tình nguyện."),
        ("WorkShifts", "EventId, StartTime, EndTime, MaxVolunteers, RequiredSkillId", "Ca làm việc."),
        ("Registrations", "EventId + UserId unique, ShiftId, Status, AttendedAt, CheckedOutAt, VolunteerHours", "Đăng ký và tham gia."),
        ("Certificates", "CertificateCode unique, UserId + EventId unique, VolunteerHours", "Chứng chỉ điện tử."),
        ("Channels/Posts/Comments/Likes/Polls", "EventId, AuthorId, PostId, UserId", "Trao đổi trong event."),
        ("SupportCampaigns/IndividualDonations", "EventId, Status, Amount, ProofImageUrl", "Ủng hộ cá nhân."),
        ("SponsorProfiles/SponsorshipProposals", "SponsorId, OrganizerId, Type, Status, ActualReceivedAmount", "Tài trợ doanh nghiệp."),
        ("Ratings/Notifications/AuditLogs", "EventId, Rater/Ratee, UserId, EntityType, Metadata", "Đánh giá, thông báo, audit."),
    ]
    add_table(doc, ["Bảng/Entity", "Trường chính", "Vai trò"], entity_rows, [2300, 4100, 2960])

    doc.add_heading("Ràng buộc dữ liệu quan trọng", level=2)
    add_bullets(
        doc,
        [
            "UserName, refresh token hash, certificate code là duy nhất.",
            "Mỗi user có tối đa một VolunteerProfile/SponsorProfile/OrganizerVerification.",
            "Mỗi volunteer chỉ có một Registration cho một Event; đăng ký lại dùng bản cũ nếu đã Cancelled.",
            "Mỗi Certificate là duy nhất theo cặp UserId + EventId.",
            "Rating duy nhất theo EventId + RaterId + RateeId.",
            "Channel shift có unique ShiftId khi không null.",
            "Like unique theo PostId + UserId; Poll vote unique theo PollId + UserId + OptionId.",
        ],
    )

    doc.add_heading("THIẾT KẾ API", level=1)
    api_rows = [
        ("POST", "/api/auth/register", "Guest", "Tạo tài khoản Volunteer/Organizer/Sponsor."),
        ("POST", "/api/auth/login", "Guest", "Đăng nhập và trả token."),
        ("GET/PUT", "/api/profile", "Authenticated", "Xem/cập nhật hồ sơ cá nhân."),
        ("POST", "/api/profile/kyc", "Volunteer", "Gửi hồ sơ KYC."),
        ("GET/POST/PUT/DELETE", "/api/events", "Public/Organizer/Admin", "Tìm kiếm, tạo, sửa, xóa có điều kiện."),
        ("PUT", "/api/events/{id}/approve|reject|cancel|complete", "Admin/Organizer", "Thay đổi trạng thái sự kiện."),
        ("POST/DELETE", "/api/events/{id}/register", "Volunteer", "Đăng ký hoặc rút đăng ký."),
        ("PUT", "/api/events/{id}/registrations/{regId}/confirm|cancel", "Organizer", "Xác nhận/hủy registration."),
        ("POST", "/api/events/{id}/registrations/{regId}/checkin", "Organizer", "Check-in QR/GPS."),
        ("POST", "/api/events/{id}/self-checkin", "Volunteer", "Tự check-in."),
        ("POST", "/api/events/{id}/registrations/{regId}/checkout", "Organizer", "Check-out và tính giờ."),
        ("GET/POST", "/api/events/{id}/support-campaigns", "Public/Organizer", "Danh sách/tạo campaign."),
        ("POST/PUT", "/api/support-campaigns/{id}/donations, /api/donations/{id}/confirm", "Volunteer/Organizer", "Ủng hộ và xác nhận donation."),
        ("POST/PUT", "/api/events/{id}/sponsorship-proposals/*", "Organizer/Sponsor", "Đề nghị, phản hồi, ghi nhận, báo cáo tài trợ."),
        ("GET", "/api/admin/monitoring/summary, /api/admin/audit-logs", "Admin", "Giám sát hệ thống và audit."),
    ]
    add_table(doc, ["Method", "Endpoint", "Role", "Mục đích"], api_rows, [900, 3300, 1500, 3660])

    doc.add_heading("THIẾT KẾ GIAO DIỆN", level=1)
    route_rows = [
        ("Public", "/", "LandingPage", "Giới thiệu hệ thống và CTA theo vai trò."),
        ("Public", "/events, /events/:id", "EventList, EventDetail", "Khám phá và xem chi tiết event."),
        ("Auth", "/login, /register", "Login, Register", "Đăng nhập/đăng ký."),
        ("Shared", "/dashboard, /notifications, /channels/:id", "Dashboard, Notifications, Channel", "Màn hình dùng chung sau đăng nhập."),
        ("Volunteer", "/profile, /activity, /achievements", "Profile, Activity, Achievements", "Hồ sơ, đăng ký/donation, chứng chỉ/huy hiệu."),
        ("Organizer", "/my-events, /events/create, /events/:id/manage", "MyEvents, EventForm, ManageEvent", "Quản lý event và vận hành hiện trường."),
        ("Sponsor", "/my-sponsorships, /sponsor/profile", "MySponsorships, SponsorProfile", "Hồ sơ và tài trợ."),
        ("Admin", "/admin/*", "Admin pages", "Duyệt, vận hành, monitoring, export."),
    ]
    add_table(doc, ["Nhóm", "Route", "Component", "Chức năng"], route_rows, [1300, 2500, 2500, 3060])

    doc.add_heading("THIẾT KẾ LUỒNG XỬ LÝ", level=1)
    add_figure(
        doc,
        DIAGRAM_DIR / "uml_sequence_approve_event.png",
        "Hình 5. Sequence Diagram cho luồng Admin duyệt sự kiện.",
    )
    add_figure(
        doc,
        DIAGRAM_DIR / "uml_state_machines.png",
        "Hình 6. State Machine Diagram cho Event, Registration, Donation và Proposal.",
    )
    flows = [
        ("Luồng tạo và duyệt sự kiện", [
            "Organizer nhập event form và gửi API POST /events.",
            "Backend tạo Event trạng thái Pending.",
            "Admin approve; EventService kiểm organizer active/verified, sinh QR, tạo Channel.",
            "Hệ thống gửi notification cho organizer; event xuất hiện public khi Approved.",
        ]),
        ("Luồng đăng ký và điểm danh", [
            "Volunteer đăng ký event; RegistrationService kiểm status, thời gian, capacity, KYC, shift.",
            "Organizer xác nhận registration; volunteer nhận notification.",
            "Volunteer/Organizer check-in bằng QR hoặc GPS trong cửa sổ hợp lệ.",
            "Organizer check-out; hệ thống tính giờ thực tế và cập nhật passport/badge.",
        ]),
        ("Luồng hoàn thành và cấp chứng chỉ", [
            "Organizer/Admin complete event Approved sau khi đủ điều kiện.",
            "Registration Pending hoặc cancel request không được tính tham gia.",
            "CertificateService cấp certificate cho volunteer attended.",
            "Guest có thể verify certificate bằng mã code.",
        ]),
        ("Luồng campaign/donation", [
            "Organizer tạo campaign Draft/Open cho event Approved.",
            "Volunteer gửi donation PendingConfirmation kèm minh chứng.",
            "Organizer confirm/reject; chỉ Confirmed được tính vào impact.",
            "Sau khi đóng campaign, organizer report sử dụng tiền.",
        ]),
        ("Luồng sponsorship proposal", [
            "Organizer request sponsor hoặc Sponsor offer event.",
            "Bên nhận accept/reject; organizer mark received khi nhận tiền.",
            "Received/Reported được tính vào impact; Pending/Accepted bị Cancelled khi event hủy.",
            "Admin finance watch xem các proposal/donation/campaign cần đối soát.",
        ]),
    ]
    for title, steps in flows:
        doc.add_heading(title, level=2)
        add_numbered(doc, steps)

    doc.add_heading("BẢO MẬT, KIỂM SOÁT LỖI VÀ AUDIT", level=1)
    add_bullets(
        doc,
        [
            "Frontend lưu token và refresh token trong localStorage, interceptor tự refresh khi gặp 401 hợp lệ.",
            "ProtectedRoute giới hạn route theo role; backend vẫn là lớp bảo vệ chính bằng [Authorize]/role checks.",
            "IsActive middleware chặn tài khoản bị khóa trên các service.",
            "Rate limiter áp dụng cho nhóm API nhạy cảm.",
            "Reject reason, request changes, delete có điều kiện và finance transition được validate ở backend.",
            "AuditLog ghi thao tác vận hành, đặc biệt thao tác admin, check-in/attendance, export và thay đổi dữ liệu nhạy cảm.",
        ],
    )

    add_nfr_design_response(doc)

    doc.add_heading("TRIỂN KHAI VÀ KIỂM THỬ", level=1)
    add_table(
        doc,
        ["Hạng mục", "Lệnh/đường dẫn", "Ghi chú"],
        [
            ("Restore backend", "dotnet restore BaseCore.sln", "Chạy tại thư mục BaseCore."),
            ("Build backend", "dotnet build BaseCore.sln --no-incremental", "Kiểm tra compile toàn solution."),
            ("Install frontend", "npm install", "Chạy tại BaseCore.WebClient."),
            ("Build frontend", "npm run build", "Kiểm tra Vite build."),
            ("Chạy local", "Auth 5002, API 5001, Gateway 5000, Web 3000", "Frontend gọi /api qua gateway."),
            ("E2E tests", "BaseCore.WebClient/tests/e2e", "Smoke, role access, public flow, business flows."),
        ],
        [1800, 3600, 3960],
    )

    doc.add_heading("PHỤ LỤC", level=1)
    doc.add_heading("Tài khoản demo", level=2)
    add_table(
        doc,
        ["Role", "Username", "Email", "Password"],
        [
            ("Admin", "admin", "admin@volunteerhub.vn", "admin123"),
            ("Organizer", "organizer", "organizer@volunteerhub.vn", "organizer123"),
            ("Sponsor", "sponsor", "sponsor@volunteerhub.vn", "sponsor123"),
            ("Volunteer", "volunteer", "volunteer@volunteerhub.vn", "volunteer123"),
        ],
        [1600, 1900, 3600, 2260],
    )
    add_detail_deep_content(doc)
    return doc


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    generate_uml_diagrams()
    docs = [
        ("VolunteerHub_SRS.docx", build_srs()),
        ("VolunteerHub_DetailDesign.docx", build_detail_design()),
    ]
    for name, doc in docs:
        apply_report_numbering(doc)
        path = OUT_DIR / name
        doc.save(path)
        print(path)


if __name__ == "__main__":
    main()
